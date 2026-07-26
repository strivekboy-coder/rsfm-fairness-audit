from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(
    r"D:\Codex\rsfm-fairness-audit\reports\advisor_geobwer_upgrade_2026_07_26"
)
FILES = (
    "GeoBWER项目汇报_2026-07-26.docx",
    "GeoBWER项目阶段性升级汇报_详细讲稿与答疑手册_2026-07-26.docx",
)
REQUIRED = (
    "reBEN",
    "27",
    "63",
    "3.3",
    "8.3",
    "4.3%",
    "24.7%",
)
FORBIDDEN = (
    "完成 reBEN 剩余",
    "顺序是先完成 reBEN",
)


for filename in FILES:
    path = ROOT / filename
    doc = Document(path)
    text = "\n".join(item.text for item in doc.paragraphs)
    text += "\n" + "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    page_breaks = 0
    for paragraph in doc.paragraphs:
        for node in paragraph._p.iter():
            if node.tag == qn("w:br") and node.get(qn("w:type")) == "page":
                page_breaks += 1

    missing = [item for item in REQUIRED if item not in text]
    stale = [item for item in FORBIDDEN if item in text]
    metrics = {
        "file": filename,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "table_rows": sum(len(table.rows) for table in doc.tables),
        "characters": len(text),
        "headings": sum(
            1 for item in doc.paragraphs if item.style.name.startswith("Heading")
        ),
        "manual_page_breaks": page_breaks,
        "replacement_characters": text.count("\ufffd"),
        "missing_required_reben_terms": missing,
        "stale_status_phrases": stale,
    }
    print(metrics)
    if text.count("\ufffd") or missing or stale:
        raise SystemExit(f"QA failed for {filename}: {metrics}")
