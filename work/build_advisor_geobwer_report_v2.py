from __future__ import annotations

from pathlib import Path
from typing import Sequence

import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports" / "advisor_geobwer_upgrade_2026_07_26"
ASSET_DIR = OUT_DIR / "assets"
REPORT_PATH = OUT_DIR / "GeoBWER项目阶段性升级汇报_导师版_重制稿_2026-07-26.docx"
SCRIPT_PATH = OUT_DIR / "GeoBWER项目阶段性升级汇报_简明讲稿_2026-07-26.docx"

NAVY = "17365D"
BLUE = "2E74B5"
MID_BLUE = "5B9BD5"
PALE_BLUE = "EAF2F8"
VERY_PALE_BLUE = "F5F8FB"
GRAY_FILL = "F2F4F7"
GRAY_LINE = "CCD5DF"
TEXT = "202A35"
MUTED = "5F6B78"
WHITE = "FFFFFF"


def font(run, size=None, bold=None, color=None, italic=None, latin="Calibri", east="Microsoft YaHei"):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def shade(element, fill):
    props = element.get_or_add_tcPr() if hasattr(element, "get_or_add_tcPr") else element
    shd = props.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        props.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_shade(paragraph, fill, left_border=BLUE):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), left_border)
    pbdr.append(left)
    ppr.append(pbdr)


def set_cell_margins(cell, top=95, bottom=95, start=120, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.find(qn("w:tcMar"))
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tcmar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    trpr.append(node)


def prevent_row_split(row):
    trpr = row._tr.get_or_add_trPr()
    trpr.append(OxmlElement("w:cantSplit"))


def configure_table(table, ratios: Sequence[float], total=9360):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [round(total * value / sum(ratios)) for value in ratios]
    widths[-1] += total - sum(widths)
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:type"), "dxa")
    tblw.set(qn("w:w"), str(total))
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:type"), "dxa")
    tblind.set(qn("w:w"), "120")
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            cell.width = Twips(widths[index])
            tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tcw.set(qn("w:type"), "dxa")
            tcw.set(qn("w:w"), str(widths[index]))
            set_cell_margins(cell)
    return widths


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def configure_doc(title):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.78)
    sec.bottom_margin = Inches(0.78)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.38)
    sec.footer_distance = Inches(0.38)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(NAVY)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(BLUE)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(5)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(NAVY)
    h3.paragraph_format.space_before = Pt(7)
    h3.paragraph_format.space_after = Pt(3)
    h3.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    r = header.add_run(title)
    font(r, size=8.5, color=MUTED)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    r = footer.add_run("RSFM / GeoFM deployment-slice reliability  ·  ")
    font(r, size=8, color=MUTED)
    add_page_field(footer)
    return doc


def title_block(doc, kicker, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(kicker.upper())
    font(r, size=9.5, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    font(r, size=25, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(subtitle)
    font(r, size=12.5, color=MUTED)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)
    ppr = rule._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    pbdr.append(bottom)
    ppr.append(pbdr)


def h1(doc, text):
    return doc.add_paragraph(text, style="Heading 1")


def h2(doc, text):
    return doc.add_paragraph(text, style="Heading 2")


def h3(doc, text):
    return doc.add_paragraph(text, style="Heading 3")


def prose(doc, text, first_bold=None, after=6, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if first_bold and text.startswith(first_bold):
        r = p.add_run(first_bold)
        font(r, size=size, bold=True, color=NAVY)
        r = p.add_run(text[len(first_bold):])
        font(r, size=size, color=TEXT)
    else:
        r = p.add_run(text)
        font(r, size=size, color=TEXT)
    return p


def note(doc, lead, text, fill=VERY_PALE_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.04)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.18
    paragraph_shade(p, fill)
    r = p.add_run(lead + " ")
    font(r, size=10.5, bold=True, color=NAVY)
    r = p.add_run(text)
    font(r, size=10.5, color=TEXT)
    return p


def formula(doc, *lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.18
    paragraph_shade(p, PALE_BLUE, MID_BLUE)
    for i, line in enumerate(lines):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        font(r, size=12, color=NAVY, latin="Cambria Math")
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    font(r, size=10.7, color=TEXT)
    return p


def table(doc, headers, rows, ratios, size=9.0):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    set_repeat_header(t.rows[0])
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade(cell._tc, PALE_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        font(r, size=size, bold=True, color=NAVY)
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        for ci, value in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri % 2:
                shade(cell._tc, "FAFBFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(str(value))
            font(r, size=size, bold=(ci == 0), color=NAVY if ci == 0 else TEXT)
    configure_table(t, ratios)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return t


def page_break(doc):
    doc.add_page_break()


def create_core_figure(path):
    fpath = Path(r"C:\Windows\Fonts\msyh.ttc")
    if fpath.exists():
        plt.rcParams["font.family"] = font_manager.FontProperties(fname=str(fpath)).get_name()
    fig, ax = plt.subplots(figsize=(12, 5.9), dpi=190)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 5.9)
    ax.axis("off")

    def box(x, y, w, h, fill, edge, title, body):
        patch = FancyBboxPatch(
            (x, y), w, h,
            boxstyle="round,pad=0.07,rounding_size=0.16",
            linewidth=1.7, facecolor=fill, edgecolor=edge,
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h * 0.69, title, ha="center", va="center",
                fontsize=13, fontweight="bold", color="#17365D")
        ax.text(x + w / 2, y + h * 0.34, body, ha="center", va="center",
                fontsize=10.7, color="#202A35", linespacing=1.35)

    box(3.25, 4.35, 5.5, 1.1, "#EAF2F8", "#2E74B5",
        "共同核心", "最差部署切片相对平均水平，多承担了多少风险？")
    box(0.45, 0.75, 5.15, 2.25, "#F5F8FB", "#8EA9C1",
        "BWER1（上次汇报）",
        "选择最差的若干“完整切片”\n再减去所有切片的平均风险")
    box(6.4, 0.75, 5.15, 2.25, "#EAF2F8", "#2E74B5",
        "GeoBWER（当前版本）",
        "精确选取最差 β 部署质量\n再减去审计测度 μ 下的平均风险")
    ax.annotate("", xy=(3.08, 2.98), xytext=(5.25, 4.34),
                arrowprops=dict(arrowstyle="-|>", color="#687789", lw=2))
    ax.annotate("", xy=(8.92, 2.98), xytext=(6.75, 4.34),
                arrowprops=dict(arrowstyle="-|>", color="#2E74B5", lw=2))
    ax.text(6, 0.22, "保留原始公平性问题；升级估计对象、部署权重、支持规则与统计认证",
            ha="center", va="center", fontsize=10.8, color="#5F6B78")
    fig.tight_layout(pad=0.2)
    fig.savefig(path, facecolor="white", bbox_inches="tight")
    plt.close(fig)


def add_figure(doc, path, width, alt, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", title)


def relax_report_spacing(doc):
    """Increase Chinese reading comfort without changing the content hierarchy."""
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name.startswith("Heading"):
            continue
        ppr = paragraph._p.get_or_add_pPr()
        shaded = ppr.find(qn("w:shd")) is not None
        paragraph.paragraph_format.line_spacing = 1.28 if shaded else 1.36
        if not shaded and style_name in {"Normal", "List Bullet"}:
            paragraph.paragraph_format.space_after = Pt(7.5)
    for table_item in doc.tables:
        for row in table_item.rows:
            for cell in row.cells:
                set_cell_margins(cell, top=115, bottom=115, start=120, end=120)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.18


def build_report():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    core_fig = ASSET_DIR / "bwer_core_v2.png"
    create_core_figure(core_fig)

    doc = configure_doc("GeoBWER 项目阶段性升级汇报")
    doc.core_properties.title = "GeoBWER项目阶段性升级汇报"
    doc.core_properties.subject = "BWER1升级、AlphaEarth和Conformal"
    doc.core_properties.author = "RSFM Fairness Audit Project"

    # P1
    title_block(
        doc,
        "阶段性研究汇报 · 2026.07",
        "GeoBWER：从尾部差异指标到 GeoFM 公平性审计框架",
        "BWER1.0 之后的方法重构、实验矩阵扩展，以及 AlphaEarth 与 Conformal 的加入",
    )
    h1(doc, "本次升级的核心结论")
    prose(
        doc,
        "项目最初的问题没有改变：GeoFM 的平均性能可能很好，但某些灾害事件、国家、类别或传感器条件可能承担明显更高的失败风险。"
        "BWER1 已经证明这种“平均值掩盖尾部”的现象值得研究；当前工作不是另起炉灶，而是把这个直觉升级为定义更精确、比较更公平、能够报告统计证据强弱的 GeoBWER。",
    )
    prose(
        doc,
        "上次汇报以后，工作新增了两条此前尚未完整呈现的证据线。第一条是 AlphaEarth 全球土地覆盖审计，它把项目扩展到 111 个国家、空间块划分和参考地图歧义；"
        "第二条是 Conformal Prediction / Conformal Risk Control，它不再只问模型“答得对不对”，还问模型给出的不确定性承诺是否在不同部署切片上被同等兑现。",
    )
    note(
        doc,
        "最简化的理解：",
        "BWER1 像一把能够发现问题的尺子；GeoBWER 保留同一刻度含义，但补上了精确尾部、部署权重、支持量、共同支持比较、空间依赖区间和无效状态，因此更接近可被其他 GeoFM 研究复用的正式测评协议。",
    )
    h2(doc, "目前项目由三个层次组成")
    bullet(doc, "一个核心风险泛函：GeoBWER = 尾部部署风险 − 平均部署风险。")
    bullet(doc, "一个任务无关审计协议：Raw、Standardised、Selective、Conformal/CRC 都进入同一泛函。")
    bullet(doc, "一个四任务验证套件：事件、地理、传感器模态和全球地图协议分别承担不同证据角色。")

    # P2
    page_break(doc)
    h1(doc, "1. 实验矩阵发生了什么变化")
    prose(
        doc,
        "BWER1 阶段已经覆盖洪水分割、地理不相交分类和多模态多标签分类，但一些比较仍是“协议感知的并列结果”："
        "模型可能使用不同波段、不同划分或不同训练预算，因此可以说明部署风险存在，却不一定能把差异归因于基础模型本身。"
        "当前矩阵的重点不是简单增加模型数量，而是为每个关键任务建立同 split、同输入条件、多个随机种子的主比较，同时把任务专用强模型保留为外部有效性参考。",
    )
    table(
        doc,
        ["任务", "BWER1 阶段", "当前正式矩阵", "升级意义"],
        [
            [
                "Sen1Floods11\n洪水分割",
                "Prithvi TL、U-Net、ResNet34-U-Net；以 S2 和事件级分析为主。",
                "TerraMind 与 ResNet34-U-Net 均运行 S1、S2、S1+S2 和 3 seeds；Prithvi TL 为外部强参考。",
                "把事件尾部风险与传感器模态机制连接起来；主比较共享官方 split。",
            ],
            [
                "fMoW-Sentinel\n单标签分类",
                "DOFA 与 13-band ResNet-50；位置不相交，但输入与训练协议并不完全对齐。",
                "DOFAv2 与 common-9-band ResNet-50，各 3 seeds；保存完整 62 维概率。",
                "可以做共同支持配对比较、排名反转检验和正式 prediction sets。",
            ],
            [
                "reBEN\n多标签分类",
                "CROMA × S1/S2/S1+S2，主要考察模态差异。",
                "CROMA、TerraMind、supervised ResNet-50 × 三种模态 × 3 seeds。",
                "形成架构×模态因子面板，判断融合收益能否跨架构复现。",
            ],
            [
                "AlphaEarth\n全球土地覆盖",
                "上次汇报尚未纳入。",
                "AlphaEarth embeddings + ML/probe；WorldCover、Dynamic World、空间块及国家×类别切片。",
                "补齐全球尺度、参考图歧义、空间转移、尺度敏感性和 coverage debt。",
            ],
        ],
        [1.05, 1.65, 2.35, 1.45],
        size=8.25,
    )
    prose(
        doc,
        "例如，旧的 Sen1 结果中 Prithvi TL 与 U-Net 都显示 Bolivia 或 Pakistan 可能成为尾部事件，但模型并不共享完全相同的训练与测试协议。"
        "新版让 TerraMind 与监督 U-Net 在同一官方划分和同一 S1/S2/融合条件下比较；如果尾部事件仍重复出现，证据更接近“数据与事件机制”，而不是“某一次模型配置偶然失败”。",
        after=4,
    )
    note(
        doc,
        "为什么没有继续堆第五个任务：",
        "现有四任务已经覆盖单标签、多标签、分割与全球地图产品，也覆盖事件、国家、类别、模态和空间块。当前论文最缺的是同协议和可认证证据，而不是更多横向案例。",
    )

    # P3
    page_break(doc)
    h1(doc, "2. BWER1 与 GeoBWER：同一个问题，更精确的定义")
    add_figure(
        doc,
        core_fig,
        6.35,
        "BWER1与GeoBWER从共同核心问题分支的概念图。",
        "BWER1与GeoBWER核心对比",
    )
    h2(doc, "两个公式")
    formula(
        doc,
        "BWER1β = 最差 k 个完整切片的平均风险 − 所有切片平均风险，k = ⌈βG⌉",
        "GeoBWERβ,μ = Tβ,μ(R) − Σg μgRg",
    )
    prose(
        doc,
        "GeoBWER 中，Tβ,μ(R) 表示审计测度 μ 下最差 β 部署质量的平均风险。μ 明确规定“每个国家等权”“按实际样本质量加权”"
        "或“按外部部署暴露量加权”；β 则规定要审计最差的多少部署质量。这样，分数不再依赖一句没有写清楚的“各组平均”。",
    )
    h2(doc, "一个最直观的例子：11 个洪水事件，β=10%")
    prose(
        doc,
        "如果 11 个事件等权，每个事件占 9.09%。BWER1 计算 ⌈0.10×11⌉=2，因此会把两个完整事件都放进尾部，实际审计 18.18% 的事件质量。"
        "GeoBWER 会取最差事件的全部 9.09%，再只取第二差事件的 0.91%，正好组成 10%。这解决了切片数较少时“10% 实际变成接近 20%”的问题。",
    )
    note(
        doc,
        "为什么有时两个版本数值仍然很接近：",
        "当切片等权且 β=k/G 恰好为整数比例时，GeoBWER 与 BWER1 数学上完全相同；fMoW 国家数很多时，临界切片只占很小质量，两者也可能近似。升级的目标是定义和证据更可靠，不是人为制造更大的分数。",
    )

    # P4
    page_break(doc)
    h1(doc, "3. BWER1 的四个隐蔽漏洞，以及新版如何修复")
    h2(doc, "漏洞一：最差切片可能只是样本少，所以碰巧很差")
    prose(
        doc,
        "假设国家 A 有 1,000 个测试样本，错误率 20%；国家 B 只有 5 个样本，其中 3 个预测错误，表面错误率是 60%。"
        "BWER1 很容易把 B 放入尾部，但 60% 可能主要是小样本波动。GeoBWER 并不简单删除 B，而是同时检查样本支持、独立 cluster 数、尾部成员稳定性，并用 cluster/spatial 同时风险带给出置信区间和单侧下界。"
        "如果证据不足，结果会明确标为 descriptive 或 inference_not_certified，而不是把一个不稳定点估计写成确定结论。",
    )
    h2(doc, "漏洞二：两个模型可能不是在同一组国家上比较")
    prose(
        doc,
        "模型 A 在 100 个国家都有有效结果，模型 B 因缺失预测只剩 80 个国家。如果各自计算 BWER，B 看起来更公平，可能只是因为最困难的 20 个国家被排除了。"
        "新版强制 common-support comparison：先冻结两个模型共同可审计的切片与独立单位，再计算配对差值区间。这样“平均性能第一、尾部公平性第二”的排名反转才有可比较含义。",
    )
    h2(doc, "漏洞三：空间样本并非相互独立")
    prose(
        doc,
        "同一卫星条带、同一事件或相邻空间块内的样本往往高度相关。把 1,000 个相邻像素当成 1,000 个独立证据，会让普通 i.i.d. bootstrap 的区间过窄。"
        "新版要求按事件、source tile、site 或空间块重采样，并使用 max-T 同时风险带传播到 GeoBWER。若正式模式缺少依赖结构字段，程序直接停止，而不是静默退化。",
    )
    h2(doc, "漏洞四：GeoBWER 为零不代表模型好")
    prose(
        doc,
        "如果所有地区错误率都是 10%，GeoBWER 接近 0；如果所有地区错误率都是 40%，GeoBWER 也可能接近 0。两者同样“均匀”，但第二个模型显然不可用。"
        "因此新版 Audit Card 强制联报平均风险、尾部风险、GeoBWER、有效部署质量和证据状态。GeoBWER 只回答风险是否分配不均，不替代准确率或总体风险。",
    )
    h2(doc, "新版如何区分“看见差距”和“证实差距”")
    prose(
        doc,
        "假设某次国家审计得到 GeoBWER=0.18，但 95% 区间为 [0, 0.42]，单侧下界 LCB=0。这个结果说明样本中观察到了明显尾部差距，"
        "但现有独立国家/site 支持还不足以排除真实差距为零；正式表述应是 descriptive evidence，而不是“已经证明不公平”。"
        "如果同样的点估计得到区间 [0.09, 0.27]、LCB=0.09，才可以说尾部超额风险在当前协议下得到统计认证。",
    )
    note(
        doc,
        "这会不会让结果显得不够震撼？",
        "一些旧的大分数可能在加入空间依赖和同时推断后变得不显著，但保留下来的结果更难被审稿人以“小样本、空间伪重复或事后挑组”推翻。新版追求的是结论强度，而不是点估计大小。",
    )

    # P5
    page_break(doc)
    h1(doc, "4. Standardised 与 Selective：原有思想被保留，但解释更严格")
    h2(doc, "Standardised-GeoBWER：控制“各地任务组成不同”")
    formula(doc, "Rstdg = Σc π*c · E[L | group=g, class=c]")
    prose(
        doc,
        "原始地理风险混合了两件事：一是某地区确实更难；二是该地区恰好包含更多困难类别。Standardised-GeoBWER 用共同参考组成 π* 重新组合每个地区的类别风险。"
        "例如，国家 A 有 80% 的机场与港口，国家 B 有 80% 的普通住宅；如果机场本来就更难，Raw-BWER 可能把类别组成差异误认为地理不公平。标准化后，两国都按照相同的类别比例比较，剩余差距更接近同类条件下的地理差异。",
    )
    prose(
        doc,
        "新版重点修复的是缺失单元。如果国家 B 根本没有“港口”样本，直接把剩余类别重新归一化会偷偷改变问题。现在 strict 版本会标记为不可识别；"
        "overlap 只在共同存在的类别上比较；partial bounds 则给出缺失单元在合理风险范围内可能导致的上下界。它们不是三个随意挑选的分数，而是三个证据层次。",
    )
    note(
        doc,
        "一个可能发生的结论变化：",
        "国家 A 的 Raw 风险为 30%，国家 B 为 15%，看起来 A 明显更差；但 A 的样本中困难类别占 80%，B 只有 20%。按同一类别组成标准化后，两国都可能变成约 20%。"
        "这时合理结论不是“新版把不公平抹掉了”，而是原始差距主要由任务组成解释。反过来，如果标准化后 A 仍为 30%、B 仍为 15%，地理差距证据反而更强。",
    )
    h2(doc, "Selective-GeoBWER：模型可以拒答，但不能用拒绝服务制造公平")
    formula(
        doc,
        "Rselg(τ) = E[L | confidence≥τ, group=g]",
        "coverageg(τ) = P(confidence≥τ | group=g)",
    )
    prose(
        doc,
        "Selective prediction 让模型只处理高置信样本。旧版主要观察“拒绝低置信样本后，准确率和 BWER 是否改善”；新版强制把每个切片的接受率一起报告。"
        "例如，模型在欧洲接受 90% 的样本，在非洲只接受 5%。即使两边接受后的错误率都只有 5%，也不能称为公平，因为非洲承担了更高的服务排除。",
    )
    note(
        doc,
        "一个关键修复：",
        "如果某个预注册切片在阈值下一个样本都不接受，其选择性风险不是 0，而是不可识别。新版不会再输出一个看似完美、实际来自完全拒答的公平分数。",
    )
    h2(doc, "怎样解读风险—覆盖曲线")
    prose(
        doc,
        "假设把置信阈值从 0.5 提高到 0.9 后，总体错误率从 20% 降到 5%。如果所有地区接受率都从约 90% 同步降到约 60%，且 GeoBWER 也下降，可以说拒答策略较均衡地改善了可靠性。"
        "如果欧洲仍接受 70%，非洲却降到 5%，那么低错误率主要是把困难地区排除在服务之外。新版会同时画出 risk、coverage 和 coverage-GeoBWER，而不是只展示一条越来越好看的准确率曲线。",
    )

    # P6
    page_break(doc)
    h1(doc, "5. Conformal：公平性审计从“预测错误”扩展到“可靠性承诺”")
    h2(doc, "最通俗的理解")
    prose(
        doc,
        "普通分类必须给出一个答案，例如“草地”。Conformal Prediction 可以给出集合 {草地，灌木地}：模型并没有突然变准，而是诚实表示这两个类别都合理。"
        "我们提前设定 90% coverage，并用独立 calibration 集决定集合需要多宽；目标是在与校准条件相似的未来样本中，真实类别大约至少 90% 的时间落在集合里。",
    )
    note(
        doc,
        "一个简单例子：",
        "模型概率为草地 0.55、灌木地 0.30、农田 0.10。普通预测输出“草地”；Conformal 可能输出 {草地，灌木地}。困难样本的集合会更大，因此 coverage 必须与平均集合大小一起解释：如果每次都把所有类别放进去，覆盖率虽高，却没有使用价值。",
    )
    h2(doc, "Conformal-GeoBWER 做了什么")
    prose(
        doc,
        "对每个测试样本定义误覆盖损失 L=1{真实标签不在预测集合中}，再计算各国家、类别或事件的误覆盖率，最后用 GeoBWER 审计哪些切片承担更多误覆盖。"
        "它回答的不是“谁的分类错误更多”，而是“当系统承诺 90% coverage 时，谁更经常得不到这个承诺”。",
    )
    prose(
        doc,
        "Conformal 与 Selective 不是替代关系。Selective 是“给一个答案或拒答”；Conformal 是“给一个可能包含多个答案的集合”。"
        "前者需要检查服务覆盖率，后者需要检查 coverage、集合大小和 coverage debt。两者都可以进入同一 GeoBWER 泛函，但风险输入不同。",
    )
    h2(doc, "同一原则在三类任务中是什么样子")
    prose(
        doc,
        "在 AlphaEarth 和 fMoW 单标签分类中，输出是真实类别可能所在的候选集合，例如 {机场，港口，工业设施}。"
        "在 reBEN 多标签任务中，一张图本来就可能同时属于农田、草地和水体，因此重点不是再造一个单标签集合，而是用 CRC 控制平均漏掉多少真实标签。"
        "在 Sen1Floods11 分割中，逐像素输出 {洪水，非洪水} 往往会形成大面积模糊集合；更可解释的做法是控制每个事件或图块的洪水漏检风险，再检查哪些事件超过目标。",
    )
    note(
        doc,
        "为什么这仍然算统一框架：",
        "三类任务不共享完全相同的集合公式，但共享同一个接口：先由任务正确的方法产生逐独立单位风险，再聚合成部署切片风险，最后由 GeoBWER 审计尾部。统一的是科学问题和审计合同，而不是强迫所有任务输出同一种对象。",
    )
    h2(doc, "AlphaEarth 已出现的信号")
    prose(
        doc,
        "AlphaEarth 在目标 90% coverage 下，空间块测试的边际 coverage 约为 88.99%，平均集合大小约为 1.449。"
        "这说明集合总体仍较精简，但空间转移已经造成轻微的总体覆盖债务。下一步更重要的问题不是只看 88.99%，而是检查哪些国家或土地类型承担了主要欠覆盖。",
    )
    note(
        doc,
        "为什么不能只看 Conformal-GeoBWER：",
        "如果所有地区的误覆盖率都是 30%，地区差距可以接近 0，但 90% coverage 承诺已经整体失败。因此报告必须同时包含总体 coverage、最差切片 coverage、集合大小、目标违反量和 Conformal-GeoBWER。",
    )

    # P7
    page_break(doc)
    h1(doc, "6. GeoConformal 论文如何融入本项目")
    prose(
        doc,
        "Lou、Luo 与 Meng 的 GeoConformal Prediction 把地理距离引入 conformal 校准：对某个测试地点而言，附近校准样本的误差权重更高，因此不同地点可以得到不同的局部阈值。"
        "它最有价值的启发是把“全球 coverage 正常，但某个局部区域持续失败”变成可观测问题。这与本项目寻找最差地理切片的目标高度一致。",
    )
    h2(doc, "为什么没有把原算法直接套到四个任务")
    prose(
        doc,
        "任意的地理核权重并不会自动继承普通 conformal 的有限样本覆盖保证。Covariate-shift weighted CP 需要权重能解释为密度比或满足加权可交换性；"
        "Localized CP 也需要专门校准。后续局部加权研究进一步表明，朴素的测试点中心 local CP 可能欠覆盖或过覆盖。"
        "因此当前设计把普通 split CP / CRC 作为正式理论锚点，把地理核版本标记为 empirical geographic localization comparator。",
    )
    prose(
        doc,
        "这并不是降低贡献。两层结构分别回答两个问题：正式锚点回答“覆盖或风险承诺是否有理论依据”；空间 comparator 回答“地理局部化能否减少最差地区的欠覆盖，以及是否以更大的预测集合为代价”。"
        "带宽只从校准坐标选择，测试标签不参与；同时报告局部有效样本量、最近校准距离、可识别比例和集合效率。",
    )
    table(
        doc,
        ["任务", "正式理论锚点", "地理局部决策"],
        [
            ["AlphaEarth", "LAC / APS / RAPS", "正式运行地理核 comparator；全球坐标与空间块最完整。"],
            ["fMoW-Sentinel", "LAC / APS / RAPS", "坐标来源和局部有效样本量门控通过后运行。"],
            ["reBEN", "多标签 CRC", "CRC 已完成；缺少可靠连续坐标，地理核版本正式标记 screened_not_run。"],
            ["Sen1Floods11", "分割 CRC", "按芯片/事件支持决定；支持不足时正式停止局部扩展。"],
        ],
        [1.35, 2.05, 3.1],
        size=8.9,
    )
    prose(
        doc,
        "因此，“screened_not_run”也可以是有意义的结果：它说明该数据集当前没有足够空间支持来识别局部不确定性，而不是程序漏做了一项分析。",
        after=3,
    )
    note(
        doc,
        "最终可证伪问题：",
        "当地理误差非平稳时，局部 conformal 能否降低最差地区的误覆盖，又不把代价转移为这些地区更大、近乎无用的预测集合？无论改善、无效还是发生代价转移，都能由 GeoBWER 的误覆盖与效率双审计揭示。",
    )

    # P8
    page_break(doc)
    h1(doc, "7. AlphaEarth：为什么它是本次新增中最重要的实验")
    prose(
        doc,
        "AlphaEarth full v2 150k 包含 156,246 个样本、111 个国家以及完整类别概率。相较前三个任务，它同时拥有全球坐标、大样本、空间块和两套土地覆盖参考产品，"
        "因此不仅能展示 GeoBWER 是否发现尾部差距，还能研究尾部差距为什么出现、支持量如何改变发现能力，以及参考地图不确定性会不会被误写成模型公平性问题。",
    )
    h2(doc, "实验设计")
    bullet(doc, "按空间块划分 train、calibration 和 test，避免相邻位置随机分割造成过度乐观。")
    bullet(doc, "审计 country、class、country×class 和 spatial block，并进行多随机种子与样本规模敏感性。")
    bullet(doc, "以 ESA WorldCover 作为主要评价参考，同时使用 Dynamic World 识别参考产品分歧与语义模糊区域。")
    bullet(doc, "从同一完整概率资产派生 Raw、Standardised、Selective、split conformal 和地理局部 comparator，避免重复推理。")
    h2(doc, "一个重要发现：参考地图不是无噪声真值")
    prose(
        doc,
        "WorldCover 与 Dynamic World 的一致率约为 50.3%。这意味着模型与 WorldCover 不一致时，不能自动把全部差异解释为“模型错误”或“模型对某国不公平”。"
        "一部分差异可能来自时间不一致、空间分辨率、土地覆盖定义或地图产品自身误差。新版因此把该风险表述为 map disagreement / evaluation-reference risk，并检查差距是否在两套参考产品一致区域仍然存在。",
    )
    note(
        doc,
        "一个具体情境：",
        "某国大量区域处于草地—灌木地过渡带。WorldCover 标为草地，Dynamic World 标为灌木地，而模型也预测灌木地。旧叙事会把它算作该国模型错误；新版会把它标记为参考图歧义，并比较“产品一致区”和“产品分歧区”的风险。如果尾部主要来自分歧区，结论应是地图协议不确定性分布不均，而不是直接指控模型偏见。",
    )
    h2(doc, "AlphaEarth 在论文中的角色")
    prose(
        doc,
        "Sen1 展示小样本事件尾部，fMoW 展示现实稀疏交叉切片，reBEN 展示传感器模态机制；AlphaEarth 则提供支持更密集的全球场景。"
        "它有机会给出非零 LCB、较完整的切片覆盖和空间 coverage debt，从而证明 GeoBWER 不只是诚实地报告“数据不足”，也能在数据支持充分时认证实质性差距。",
    )

    # P9
    page_break(doc)
    h1(doc, "8. reBEN：升级后的框架已经发现了什么")
    prose(
        doc,
        "reBEN 正式面板已经完成：CROMA、TerraMind 与监督 ResNet-50，在 S1、S2、S1+S2 三种输入和三个随机种子下共 27 次运行。"
        "下表为三 seed 均值，风险与 GeoBWER 均为越低越好。它第一次让“平均表现、国家尾部、风险控制和服务覆盖”在同一受控实验中同时可比较。",
    )
    table(
        doc,
        ["模型管线", "模态", "样本风险", "尾部风险", "GeoBWER"],
        [
            ["ResNet-50", "S1", "0.1281", "0.1931", "0.0488"],
            ["ResNet-50", "S2", "0.0927", "0.1409", "0.0400"],
            ["ResNet-50", "S1+S2", "0.0876", "0.1325", "0.0371"],
            ["CROMA", "S1", "0.1065", "0.1583", "0.0391"],
            ["CROMA", "S2", "0.0925", "0.1380", "0.0375"],
            ["CROMA", "S1+S2", "0.0870", "0.1315", "0.0366"],
            ["TerraMind", "S1", "0.1121", "0.1718", "0.0442"],
            ["TerraMind", "S2", "0.0804", "0.1193", "0.0310"],
            ["TerraMind", "S1+S2", "0.0801", "0.1197", "0.0311"],
        ],
        [1.55, 0.85, 1.35, 1.35, 1.4],
        size=8.55,
    )
    h2(doc, "这张表最值得讲的三个结果")
    bullet(
        doc,
        "光学信息的可靠性优势非常稳定：三个模型、三个 seeds、七个有支持国家组成的 63 个配对比较中，S1 风险均高于 S2；S1 与融合输入的 63 个比较也全部同方向。",
    )
    bullet(
        doc,
        "平均更好不保证尾部同步改善：TerraMind 从 S2 融合到 S1+S2 后，平均风险由 0.0804 微降至 0.0801，但尾部风险由 0.1193 微升至 0.1197；九组排名中有四组出现平均风险与 GeoBWER 排名不一致。",
    )
    bullet(
        doc,
        "总体承诺会掩盖地理债务：CRC 的总体风险约为 9.77%–10.70%，接近 10% 目标，但最差国家仍额外承担 3.31–8.32 个百分点的风险；Selective 的总体接受率约 50%，最差国家却只有 4.3%–24.7%。",
    )
    note(
        doc,
        "怎样证明这不是旧 BWER 换名字：",
        "旧版主要能给出尾部点估计；新版把同一现象拆成可比较的平均风险、尾部风险、GeoBWER、目标违反和服务覆盖，并区分描述性点估计与可认证差异。"
        "当前国家级配对改善证据很强；跨模型 GeoBWER 排名仍应称为描述性，因为最保守的同时区间尚不能排除零。",
    )
    prose(
        doc,
        "支持边界：测试集有 119,825 个样本，但正式空间推断依赖 45 个 source-tile clusters。主要点估计覆盖七个有支持国家，覆盖 97.7% 的样本、但只占等国家部署测度的 70%；因此结论应称为“七国支持全集 GeoBWER”，不能写成完整十国认证。",
        after=3,
        size=9.0,
    )

    # P10
    page_break(doc)
    h1(doc, "9. 四个实验如何组成一个统一的科学故事")
    h2(doc, "Sen1Floods11：同一种灾害任务，不同事件承担的风险并不相同")
    prose(
        doc,
        "BWER1 阶段已经发现 Prithvi TL、U-Net 和 ResNet34-U-Net 的平均 IoU 与事件尾部风险并不完全一致，Bolivia、Pakistan 等事件在部分 learned models 中重复出现。"
        "新版不假设这些旧尾部一定成立，而是用同 split 的 TerraMind / U-Net 三模态面板重新检验：增加 S1 或融合 S1+S2，究竟降低平均风险、降低尾部风险，还是只改变尾部事件身份。",
    )
    h2(doc, "fMoW-Sentinel：平均性能排名不等于地理尾部排名")
    prose(
        doc,
        "fMoW 的价值在于 62 类、位置不相交和国家支持极不均衡。旧结果提示 RSFM 与监督基线可能出现平均准确率和尾部公平性的排名分离。"
        "新版用 DOFAv2 与 common-9-band ResNet-50、三个随机种子和共同支持配对区间，判断这种排名反转能否从点估计现象升级为正式结论。"
        "Country×Class 仍保留为 supported-cell exploratory analysis，不冒充完整的全球交叉公平审计。",
    )
    h2(doc, "reBEN：融合模态是否真的缓解尾部风险")
    prose(
        doc,
        "完成的 27-run 因子面板显示，S2 相对 S1 的改善跨三个模型、三个 seeds 和全部有支持国家稳定成立；但融合并非总能同时改善平均与尾部。"
        "CROMA 在 S1 条件最好，TerraMind 在 S2 与融合条件整体最好，而 TerraMind 融合还出现“平均略好、尾部略差”。"
        "因此结论不是笼统的“融合一定更公平”，而是光学信息优势稳定、融合收益具有模型与风险层级依赖性。",
    )
    h2(doc, "AlphaEarth：全球平均可靠不等于局部覆盖承诺被公平兑现")
    prose(
        doc,
        "AlphaEarth 把项目从任务性能扩展到全球地图协议与不确定性。它检验国家、类别、空间块和参考图歧义如何共同塑造尾部风险，并用 Conformal-GeoBWER 追踪 coverage debt。"
        "四个任务因而不是四个松散案例，而是对同一问题的四种压力测试：小 G 事件、稀疏地理交叉、模态机制和大规模全球空间转移。",
    )
    note(
        doc,
        "整篇论文最希望形成的经验结论：",
        "一个模型可以在平均性能上领先，却在最差部署切片、服务接受率或不确定性覆盖上落后；只有把总体风险、尾部风险和可靠性承诺共同报告，才能判断 GeoFM 是否适合真实部署。",
    )

    # P11
    page_break(doc)
    h1(doc, "10. 当前贡献、尚未完成的证据与下一步")
    h2(doc, "目前已经形成的贡献")
    prose(
        doc,
        "第一，提出并实现 GeoBWER：一个模型无关、任务感知的尾部超额风险泛函，能够显式定义部署测度 μ 和尾部质量 β。"
        "第二，把支持诊断、strict standardisation、共同支持、空间/cluster 推断、同时风险带、LCB 和无效状态纳入同一协议。"
        "第三，通过四任务验证套件统一 Raw、Standardised、Selective、Conformal 与 CRC；reBEN 27-run 已给出首个完整架构×模态证据，AlphaEarth 则提供全球地图歧义和空间 coverage debt 分析。",
    )
    h2(doc, "目前不应提前声称的内容")
    prose(
        doc,
        "GeoBWER 的公式和协议已经固定，但最终论文证据仍取决于所有正式面板完成后的共同支持、随机种子稳定性和置信区间。"
        "不是所有观察到的正 GeoBWER 都已经得到非零 LCB；稀疏 Country×Class 只能作为有支持单元上的探索性结果；地理核 conformal 也不能表述为未经证明的点态有限样本保证。",
    )
    h2(doc, "最短且不重复的完成顺序")
    prose(doc, "1. 冻结 reBEN 27-run 主表、β-profile、CRC 与 Selective 结论；连续坐标 GeoConformal 因缺少坐标暂不运行。", first_bold="1.")
    prose(doc, "2. 完成 fMoW common-9-band ResNet-50，与 DOFAv2 做共同支持配对差值和排名反转检验。", first_bold="2.")
    prose(doc, "3. 完成 Sen1 TerraMind / U-Net 三模态面板，并把 Prithvi TL 放在外部有效性层。", first_bold="3.")
    prose(doc, "4. 冻结 AlphaEarth canonical lineage，完成 split CP 与地理局部 comparator 的 coverage—efficiency 双审计。", first_bold="4.")
    prose(doc, "5. 最后只从逐样本正式表生成 support frontier、主文图表和跨任务 Audit Cards，不再重复 GPU 推理。", first_bold="5.")
    note(
        doc,
        "最终定位：",
        "这项工作不是把 CVaR 或某个公平性公式简单移植到遥感，而是把“GeoFM 中谁算一个部署群体、如何定义目标部署质量、空间相关下如何推断、支持不足时如何拒绝伪结论、跨任务如何统一风险输入”做成一套可复用协议。"
        "论文完成后，第三方模型只需提供逐样本预测与部署元数据，即可生成 GeoBWER Audit Card。",
    )
    prose(
        doc,
        "最核心的论文叙事可以收束为：平均性能并不能回答 GeoFM 是否可靠地服务所有部署环境。GeoBWER 衡量最差 β 部署质量相对平均部署风险承担的额外风险，"
        "并通过支持、空间依赖和不确定性认证，把这一差距从直观现象升级为可复核的公平性审计证据。",
        first_bold="最核心的论文叙事可以收束为：",
        after=2,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(
        "方法来源（精选）：GeoConformal Prediction (Lou et al., 2025)；Conformal Prediction Under Covariate Shift "
        "(Tibshirani et al., 2019)；Localized Conformal Prediction (Guan, 2023)；Conformal Risk Control "
        "(Angelopoulos et al., 2024)。"
    )
    font(r, size=8.6, color=MUTED)

    relax_report_spacing(doc)
    doc.save(REPORT_PATH)
    return REPORT_PATH


def build_script():
    doc = configure_doc("GeoBWER 阶段性升级汇报 · 简明讲稿")
    doc.core_properties.title = "GeoBWER阶段性升级汇报简明讲稿"
    title_block(
        doc,
        "配套讲稿",
        "GeoBWER 项目阶段性升级汇报",
        "按新版报告 P1–P11 编排；用于会议共享时口头补充，不需要逐字照读",
    )
    script = [
        (
            "P1｜升级总览",
            "上次汇报时，核心工作是用 BWER1 证明平均性能会掩盖最差事件、国家或模态。"
            "这次没有改变方向，而是把 BWER1 升级为 GeoBWER：核心仍然是尾部风险减平均风险，但现在尾部定义更精确，"
            "并补上部署权重、支持量、空间相关区间和无效状态。新增的两条主线是 AlphaEarth 全球审计和 Conformal 不确定性审计。",
        ),
        (
            "P2｜实验矩阵",
            "这一页重点不是模型变多，而是比较变得更公平。Sen1、fMoW 和 reBEN 都增加了同 split、同模态或同波段的监督基线和三个随机种子。"
            "Prithvi 没有被删除，而是从主受控比较调整为任务专用强参考。AlphaEarth 是上次没有汇报的新全球实验。",
        ),
        (
            "P3｜BWER1 与 GeoBWER",
            "两个版本问的是同一个问题：最差部署切片比平均水平多承担多少风险。区别是 BWER1 只能取完整切片。"
            "例如 11 个事件审计最差 10%，BWER1 会取两个事件，实际变成 18.18%；GeoBWER 取最差事件全部，再取第二差事件的一小部分，精确组成 10%。",
        ),
        (
            "P4｜为什么必须升级",
            "这里可以强调四个以前不容易发现的问题：小样本国家可能偶然成为最差；两个模型可能在不同国家集合上计算；空间样本不是独立样本；"
            "GeoBWER 为零也不代表模型本身好。新版分别用支持和 LCB、共同支持配对、cluster/spatial 推断，以及平均风险和尾部风险联报解决。",
        ),
        (
            "P5｜Standardised 与 Selective",
            "Standardised 用来区分类别组成差异和真正的同类地理差异。比如一个国家机场很多，Raw 风险高不一定是地理不公平。"
            "新版遇到国家×类别缺失时不会偷偷重归一化，而是报告不可识别、共同重叠或区间。Selective 则要求把接受率一起报告，不能靠拒绝某些地区的大多数样本制造低风险。",
        ),
        (
            "P6｜Conformal",
            "普通分类给一个答案；Conformal 可以给一个集合，例如草地和灌木地。它不是让模型变准，而是让模型诚实表达不确定性。"
            "我们用独立校准集确定达到 90% coverage 所需的集合宽度，再用 GeoBWER 检查哪些国家或类别更经常得不到这个覆盖承诺。"
            "必须同时看 coverage 和集合大小，否则把全部类别都放进去也会显得很好。",
        ),
        (
            "P7｜GeoConformal",
            "导师提到的 GeoConformal 是按地理距离给附近校准样本更高权重，让不同地点有局部阈值。这个思想和项目非常契合。"
            "但任意地理核并不自动具有普通 conformal 的有限样本保证，所以我们把 split CP 或 CRC 保留为正式锚点，"
            "把地理核版本作为经验空间 comparator。它重点检验局部 coverage 是否改善，以及代价是不是集合变得更大。",
        ),
        (
            "P8｜AlphaEarth",
            "AlphaEarth 是这次新增中最完整的全球实验，有 15.6 万样本和 111 个国家。它不仅能测国家和类别尾部，还能分析 WorldCover 与 Dynamic World 的分歧。"
            "两套参考图一致率只有约 50.3%，所以不能把所有不一致都归因于模型偏见。新版会区分模型—参考图分歧和参考产品歧义。",
        ),
        (
            "P9｜reBEN 已完成结果",
            "reBEN 的 27 条正式路线已经完成。最强结果是 S2 相对 S1 的优势在三个模型、三个随机种子和七个有支持国家的 63 个比较中全部成立。"
            "同时，TerraMind 融合出现平均风险略降但尾部风险略升，九组排名中四组平均风险与 GeoBWER 次序不同。"
            "CRC 总体接近 10% 目标，最差国家仍多承担 3.3 到 8.3 个百分点；Selective 总体接受约一半样本，最差国家只接受 4.3% 到 24.7%。",
        ),
        (
            "P10｜统一故事",
            "四个任务各自承担不同压力测试：Sen1 是小样本事件和模态；fMoW 是地理泛化与排名反转；reBEN 是架构×模态机制；"
            "AlphaEarth 是全球空间转移、地图歧义和 coverage debt。它们共同说明平均性能、尾部风险、拒答覆盖和不确定性覆盖可能给出不同的模型判断。",
        ),
        (
            "P11｜当前状态与下一步",
            "方法公式、协议和主要代码已经完成，reBEN 正式面板也已经封存。"
            "下一步是完成 fMoW baseline 与共同支持比较、Sen1 三模态和 AlphaEarth 的 split conformal/空间 comparator，最后统一做 CPU 后处理和 Audit Cards。"
            "最终目标仍然是做一套其他 GeoFM 论文能够直接复用的公平性审计指标与工具。",
        ),
    ]
    for index, (title, body) in enumerate(script):
        if index in {3, 6, 8}:
            page_break(doc)
        h2(doc, title)
        prose(doc, body, after=8, size=10.8)
    doc.save(SCRIPT_PATH)
    return SCRIPT_PATH


if __name__ == "__main__":
    print(build_report())
