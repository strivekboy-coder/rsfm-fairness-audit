from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports" / "advisor_geobwer_upgrade_2026_07_25"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "GeoBWER项目阶段性升级汇报_导师版_2026-07-25.docx"

BLUE = "1F4E79"
BLUE_2 = "2E75B6"
PALE_BLUE = "EAF2F8"
PALE_CYAN = "E8F6F3"
TEAL = "147D73"
GREEN = "2E7D32"
PALE_GREEN = "EAF4EA"
ORANGE = "C55A11"
PALE_ORANGE = "FCE4D6"
RED = "B03A2E"
PALE_RED = "FDEDEC"
GOLD = "B8860B"
GRAY = "5B6573"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E0E7"
DARK = "1F2933"
WHITE = "FFFFFF"


def set_run_font(run, east_asia="Microsoft YaHei", latin="Calibri", size=None, bold=None, color=None):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs.get(edge)
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_paragraph(paragraph, fill: str, border_color: str | None = None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border_color:
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "6")
        left.set(qn("w:color"), border_color)
        p_bdr.append(left)
        p_pr.append(p_bdr)


def keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement("w:keepNext"))


def prevent_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_field(paragraph, field_code: str):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE_2)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(r_fonts)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_title(doc, text: str, subtitle: str | None = None):
    p = doc.add_paragraph()
    p.style = "Heading 1"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=22, bold=True, color=BLUE)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(14)
        run2 = p2.add_run(subtitle)
        set_run_font(run2, size=11.5, color=GRAY)


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    return p


def add_body(doc, text: str, bold_prefix: str | None = None, color=DARK, space_after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=10.5, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2, size=10.5, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=color)
    return p


def add_bullets(doc, items: Iterable[str], level=0, color=DARK, size=10.2):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, size=size, color=color)


def add_callout(doc, title: str, body: str, fill=PALE_BLUE, accent=BLUE_2, compact=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5 if compact else 8)
    p.paragraph_format.line_spacing = 1.05
    shade_paragraph(p, fill, accent)
    r1 = p.add_run(title + "  ")
    set_run_font(r1, size=10.3 if compact else 10.8, bold=True, color=accent)
    r2 = p.add_run(body)
    set_run_font(r2, size=9.8 if compact else 10.2, color=DARK)
    return p


def add_formula_block(doc, lines: Sequence[str], title: str | None = None, fill="F7F9FC"):
    if title:
        p0 = doc.add_paragraph()
        p0.paragraph_format.space_after = Pt(3)
        r0 = p0.add_run(title)
        set_run_font(r0, size=10.5, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    shade_paragraph(p, fill)
    for i, line in enumerate(lines):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, east_asia="Microsoft YaHei", latin="Cambria Math", size=12.2, color=DARK)
    return p


def add_table(
    doc,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths: Sequence[float] | None = None,
    font_size=8.8,
    header_fill=BLUE,
    first_col_fill: str | None = None,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    target_twips = 9720
    raw_widths = list(widths) if widths else [1.0] * len(headers)
    raw_total = sum(raw_widths)
    col_twips = [round(target_twips * w / raw_total) for w in raw_widths]
    col_twips[-1] += target_twips - sum(col_twips)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(target_twips))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "90")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_twips in col_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width_twips))
        grid.append(grid_col)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=font_size, bold=True, color=WHITE)
        cell.width = Twips(col_twips[i])
        tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
        if tc_w is not None:
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(col_twips[i]))
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        prevent_split(row)
        for ci, value in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            cell.width = Twips(col_twips[ci])
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is not None:
                tc_w.set(qn("w:type"), "dxa")
                tc_w.set(qn("w:w"), str(col_twips[ci]))
            if first_col_fill and ci == 0:
                set_cell_shading(cell, first_col_fill)
            elif ri % 2 == 1:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.03
            if ci == 0:
                r = p.add_run(str(value))
                set_run_font(r, size=font_size, bold=True, color=BLUE)
            else:
                r = p.add_run(str(value))
                set_run_font(r, size=font_size, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def page_break(doc):
    doc.add_page_break()


def create_core_diagram(path: Path):
    font_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    if font_path.exists():
        prop = font_manager.FontProperties(fname=str(font_path))
        plt.rcParams["font.family"] = prop.get_name()
    fig, ax = plt.subplots(figsize=(12, 6.2), dpi=180)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6.2)
    ax.axis("off")

    def box(x, y, w, h, fill, edge, title, body, title_color="#1F4E79"):
        patch = FancyBboxPatch(
            (x, y), w, h,
            boxstyle="round,pad=0.08,rounding_size=0.18",
            linewidth=1.8, facecolor=fill, edgecolor=edge,
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h * 0.67, title, ha="center", va="center",
                fontsize=14, fontweight="bold", color=title_color)
        ax.text(x + w / 2, y + h * 0.34, body, ha="center", va="center",
                fontsize=11.5, color="#1F2933", linespacing=1.35)

    box(3.2, 4.35, 5.6, 1.25, "#EAF2F8", "#2E75B6",
        "共同的核心问题", "表现最差的部署切片，是否比平均切片承担更多风险？")
    box(0.5, 0.75, 5.1, 2.15, "#F7F9FC", "#8DA9C4",
        "BWER1（上次汇报）",
        "最差若干“完整切片”的平均风险\n−\n所有切片平均风险")
    box(6.4, 0.75, 5.1, 2.15, "#E8F6F3", "#147D73",
        "GeoBWER（当前升级）",
        "精确最差 β 部署质量的尾部风险\n−\n审计测度下的平均风险",
        title_color="#147D73")
    ax.annotate("", xy=(3.05, 2.85), xytext=(5.25, 4.34),
                arrowprops=dict(arrowstyle="-|>", color="#5B6573", lw=2.1))
    ax.annotate("", xy=(8.95, 2.85), xytext=(6.75, 4.34),
                arrowprops=dict(arrowstyle="-|>", color="#147D73", lw=2.1))
    ax.text(6, 0.18, "核心思想保留；估计对象、权重、统计推断与可复用协议全面升级",
            ha="center", va="center", fontsize=11, color="#5B6573")
    fig.tight_layout(pad=0.3)
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def create_conformal_diagram(path: Path):
    font_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    if font_path.exists():
        prop = font_manager.FontProperties(fname=str(font_path))
        plt.rcParams["font.family"] = prop.get_name()
    fig, ax = plt.subplots(figsize=(12, 4.4), dpi=180)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 4.4)
    ax.axis("off")
    nodes = [
        (0.25, "模型概率", "草地 .55\n灌木 .30\n农田 .10"),
        (3.15, "独立校准集", "决定达到 90%\n覆盖率所需阈值"),
        (6.05, "预测集合", "{草地，灌木}\n而非强行单选"),
        (8.95, "GeoBWER 审计", "哪些国家/类别\n承担更多误覆盖？"),
    ]
    colors = [("#F7F9FC", "#8DA9C4"), ("#EAF2F8", "#2E75B6"),
              ("#E8F6F3", "#147D73"), ("#FCE4D6", "#C55A11")]
    for idx, (x, title, body) in enumerate(nodes):
        fill, edge = colors[idx]
        p = FancyBboxPatch((x, 1.05), 2.35, 2.3,
                           boxstyle="round,pad=0.06,rounding_size=0.15",
                           linewidth=1.6, facecolor=fill, edgecolor=edge)
        ax.add_patch(p)
        ax.text(x + 1.175, 2.78, title, ha="center", va="center",
                fontsize=12.2, fontweight="bold", color=edge)
        ax.text(x + 1.175, 1.86, body, ha="center", va="center",
                fontsize=10.5, color="#1F2933", linespacing=1.35)
        if idx < len(nodes) - 1:
            ax.annotate("", xy=(x + 2.86, 2.2), xytext=(x + 2.38, 2.2),
                        arrowprops=dict(arrowstyle="-|>", color="#5B6573", lw=1.8))
    ax.text(6, 0.38, "必须联报：总体 coverage · 最差切片 coverage · 平均集合大小 · coverage debt · Conformal-GeoBWER",
            ha="center", va="center", fontsize=10.8, color="#5B6573")
    fig.tight_layout(pad=0.2)
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def configure_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)
    sec.header_distance = Inches(0.28)
    sec.footer_distance = Inches(0.28)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, color in ((1, 16, BLUE), (2, 13, BLUE_2), (3, 11.5, TEAL)):
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if level == 1 else 7)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.2)

    header = sec.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("GeoBWER 项目阶段性升级汇报  ·  导师版")
    set_run_font(r, size=8.2, color=GRAY)
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_after = Pt(0)
    r1 = fp.add_run("RSFM / GeoFM Deployment-Slice Reliability   |   ")
    set_run_font(r1, size=8, color=GRAY)
    add_field(fp, "PAGE")
    return doc


def build_report():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    core_diagram = ASSET_DIR / "bwer1_geobwer_core.png"
    conformal_diagram = ASSET_DIR / "conformal_geobwer_pipeline.png"
    create_core_diagram(core_diagram)
    create_conformal_diagram(conformal_diagram)

    doc = configure_document()
    core = doc.core_properties
    core.title = "GeoBWER项目阶段性升级汇报"
    core.subject = "BWER1到GeoBWER、AlphaEarth与Conformal扩展"
    core.author = "RSFM Fairness Audit Project"
    core.keywords = "GeoBWER, GeoFM, RSFM, fairness, conformal prediction, AlphaEarth"

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("GeoBWER")
    set_run_font(r, size=34, bold=True, color=BLUE)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    r = p2.add_run("从“发现尾部差异”到“可复用的 GeoFM 公平性审计协议”")
    set_run_font(r, size=19, bold=True, color=TEAL)
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(22)
    r = p3.add_run("BWER1.0 之后的核心方法升级、实验矩阵扩展与不确定性设计")
    set_run_font(r, size=12, color=GRAY)

    line = doc.add_paragraph()
    line.paragraph_format.space_after = Pt(18)
    shade_paragraph(line, BLUE)
    line.add_run(" ")

    add_callout(
        doc,
        "一句话结论",
        "项目没有偏离最初的公平性指标目标：保留“最差部署切片相对平均水平的额外风险”这一核心构念，"
        "把原来的整数切片启发式升级为精确、加权、带支持诊断和统计认证的 GeoBWER。",
        fill=PALE_BLUE,
        accent=BLUE,
    )
    add_table(
        doc,
        ["上次汇报后新增的三条主线", "本次升级带来的科学意义"],
        [
            ["GeoBWER 1.1", "从一个可运行分数，升级为明确 estimand、可比较、可给置信结论的审计协议。"],
            ["AlphaEarth 全球审计", "补齐全球空间块、国家×类别、参考地图歧义与尺度敏感性证据。"],
            ["Conformal / CRC / GeoConformal", "从“模型是否答对”扩展到“模型能否诚实表达不确定性，以及哪些切片得不到可靠覆盖”。"],
        ],
        widths=[2.05, 4.7],
        font_size=9.2,
        header_fill=TEAL,
        first_col_fill=PALE_CYAN,
    )
    add_body(doc, "汇报日期：2026 年 7 月 25 日", color=GRAY, space_after=2)
    add_body(doc, "项目：RSFM / GeoFM Deployment-Slice Reliability Audit", color=GRAY, space_after=0)

    # Page 2
    page_break(doc)
    add_title(doc, "1. 这段时间真正完成了什么", "不是简单增加实验数量，而是重构了“公平性风险如何定义、比较与认证”")
    add_table(
        doc,
        ["层次", "BWER1 阶段", "当前 GeoBWER 阶段"],
        [
            ["科学问题", "平均性能之外，最差部署切片是否显著更差？", "同一问题，但明确“最差 β 部署质量”以及目标部署测度。"],
            ["指标定义", "取最差若干完整切片，再减去切片平均。", "精确取尾部 β 质量；临界切片只取所需部分；支持不同部署权重。"],
            ["证据强度", "以点估计和诊断性结果为主。", "区分 descriptive、identified、certified；报告 CI、LCB 与无效状态。"],
            ["跨模型比较", "模型可能因有效切片不同而不可直接比较。", "强制共同支持、配对差值和同协议比较。"],
            ["不确定性", "Selective-BWER 为主。", "Selective + split CP + CRC + 地理局部 comparator，统一进入 GeoBWER。"],
            ["实验结构", "四个任务较像并列案例。", "四任务分别承担事件、地理、模态和全球地图协议四类部署轴。"],
        ],
        widths=[1.1, 2.45, 3.2],
        font_size=8.65,
    )
    add_callout(
        doc,
        "核心变化",
        "GeoBWER 不是为了让数值“更大”或“更漂亮”。它的价值是把“观察到差距”与“有限样本已经能够确认差距”分开。",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )
    add_heading(doc, "导师汇报时最值得强调的三句话", 2)
    add_bullets(
        doc,
        [
            "原始 BWER 的科学直觉被保留，而且成为最终方法的嵌入特例。",
            "真正的方法创新不只是一条公式，而是“尾部风险泛函 + 部署测度 + 支持规则 + 空间依赖推断 + 版本化审计协议”。",
            "AlphaEarth 和 Conformal 不是额外拼接：前者把审计扩展到全球地图部署，后者把公平性问题扩展到可靠性承诺是否被各地公平兑现。",
        ],
    )
    add_callout(doc, "一句话讲法", "以前是“发现哪些地方差”；现在是“明确差多少、对谁而言、证据是否足够、不同模型能否公平比较”。",
                fill=PALE_GREEN, accent=GREEN, compact=True)

    # Page 3 matrix
    page_break(doc)
    add_title(doc, "2. 实验矩阵：从并列案例到四任务验证套件", "同一指标跨任务复用，但每个任务承担不同的部署可靠性问题")
    add_table(
        doc,
        ["任务", "上次汇报时的主要模型/角色", "当前冻结模型矩阵", "改动后的优势"],
        [
            [
                "Sen1Floods11\n洪水分割",
                "Prithvi TL；U-Net / ResNet34-U-Net；以 S2 与事件级 BWER 为主。",
                "TerraMind × S1/S2/S1+S2 × 3 seeds；ResNet34-U-Net 同协议三模态；Prithvi TL 作为任务专用外部强参考。",
                "把“事件尾部风险”与“传感器模态是否缓解尾部失败”连接起来；主比较共享 split、模态和训练预算。",
            ],
            [
                "fMoW-Sentinel\n单标签分类",
                "DOFA 与 13-band ResNet-50；位置不相交；模型输入和训练协议并不完全对齐。",
                "DOFAv2 × 3 probe seeds；common-9-band ResNet-50 × 3 seeds；保存完整 62 维概率。",
                "能够做同支持、同波段的 FM–baseline 配对比较；可正式运行 prediction sets 与排名反转检验。",
            ],
            [
                "reBEN\n多标签分类",
                "CROMA × S1/S2/S1+S2；主要展示传感器模态差异。",
                "CROMA、TerraMind、supervised ResNet-50 × S1/S2/S1+S2 × 3 seeds。",
                "形成“架构 × 模态”因子面板；可判断融合收益是否跨架构复现，而非单模型偶然现象。",
            ],
            [
                "AlphaEarth\n全球土地覆盖",
                "上次汇报尚未纳入。",
                "AlphaEarth embeddings + ML/probe；WorldCover 评价参考；Dynamic World 歧义诊断；空间块、国家、类别及交叉切片。",
                "补齐全球尺度、参考地图不确定性、空间转移、规模敏感性和 conformal coverage debt。",
            ],
        ],
        widths=[1.05, 1.65, 2.4, 2.0],
        font_size=7.85,
        header_fill=BLUE,
    )
    add_callout(
        doc,
        "为什么不再盲目增加第五个任务？",
        "现有四任务已覆盖单标签、多标签、像素分割与全球地图协议，也覆盖事件、地理、模态和空间尺度。"
        "当前更需要的是同协议基线、多个随机种子和可认证推断，而不是继续堆数据集。",
        fill=LIGHT_GRAY,
        accent=GRAY,
    )

    # Page 4 core diagram
    page_break(doc)
    add_title(doc, "3. 核心思想没有变：最差切片风险 − 平均风险", "BWER1 是直观起点，GeoBWER 是其精确定义与统计化升级")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    shape = r.add_picture(str(core_diagram), width=Inches(6.75))
    shape._inline.docPr.set(
        "descr",
        "BWER1 与 GeoBWER 从共同核心问题分支出的概念图：前者使用最差若干完整切片，后者使用精确最差 beta 部署质量。",
    )
    shape._inline.docPr.set("title", "BWER1 与 GeoBWER 核心思想对比")
    add_callout(
        doc,
        "最重要的理解",
        "GeoBWER 并没有推翻项目最初的科学问题；它推翻的是“用 ceil(βG) 个完整切片近似尾部 β 质量”这一实现方式。",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    add_body(
        doc,
        "其中“风险”不是固定等于分类错误率：分割可用每个事件的 IoU 损失，多标签可用漏标风险，"
        "Selective 可用接受样本上的损失，Conformal 可用误覆盖损失。因而同一个泛函可以跨任务复用。",
        space_after=4,
    )
    add_callout(doc, "一句话讲法", "BWER1 像用整块积木拼出 10%；GeoBWER 允许切开最后一块，精确得到 10%。",
                fill=PALE_ORANGE, accent=ORANGE, compact=True)

    # Page 5 formula
    page_break(doc)
    add_title(doc, "4. 两个版本的公式与一个直观例子", "数学升级发生在“尾部如何选、平均如何定义、证据如何认证”")
    add_heading(doc, "BWER1：最差若干完整切片", 2)
    add_formula_block(
        doc,
        [
            "k = ⌈βG⌉",
            "BWER1β = (1/k) Σⱼ₌₁ᵏ R(ⱼ)  −  (1/G) Σg₌₁ᴳ Rg",
        ],
    )
    add_body(doc, "R(1) ≥ R(2) ≥ … 为按风险从高到低排序的切片；G 为切片数。它直观，但当 βG 不是整数时会多取完整切片。")
    add_heading(doc, "GeoBWER：精确最差 β 部署质量", 2)
    add_formula_block(
        doc,
        [
            "平均风险：mμ(R) = Σg μg Rg",
            "尾部风险：Tβ,μ(R) = (1/β) ∫₁₋β¹ F⁻¹R,μ(u) du",
            "GeoBWERβ,μ(R) = Tβ,μ(R) − mμ(R)",
        ],
        fill="E8F6F3",
    )
    add_body(
        doc,
        "μ 是审计测度：可令国家等权、按观测样本质量加权，或在有权威依据时使用外部部署权重。"
        "β=0.10 表示审计最差 10% 的部署质量，而不是机械选择若干个组。",
    )
    add_table(
        doc,
        ["11 个洪水事件，β=10%", "实际被审计的尾部"],
        [
            ["BWER1", "⌈1.1⌉=2 个完整事件，即 18.18% 的事件质量。"],
            ["GeoBWER", "最差事件全部 9.09% + 第二差事件 0.91%，精确组成 10%。"],
        ],
        widths=[1.8, 4.9],
        font_size=9.4,
        header_fill=TEAL,
        first_col_fill=PALE_CYAN,
    )
    add_callout(
        doc,
        "嵌入性质",
        "当切片等权且 β=k/G 时，GeoBWER 与 BWER1 完全相同。因此部分数据上两者数值接近，是理论预期，不是升级失效。",
        fill=PALE_BLUE,
        accent=BLUE,
        compact=True,
    )

    # Page 6 advantages
    page_break(doc)
    add_title(doc, "5. BWER1 暴露的问题，GeoBWER 如何逐项解决", "改进目标不是复杂化，而是让分数能被正确解释和复用")
    add_table(
        doc,
        ["BWER1 阶段的问题", "GeoBWER 1.1 的对应解决", "科学价值"],
        [
            ["用整数个完整切片近似 β 尾部，随切片数跳变。", "精确 fractional boundary mass；连续 β-profile（5%、10%、20%、30%）。", "不同任务、不同切片数量之间更可比。"],
            ["默认切片等权，部署人群含义不够明确。", "显式审计测度 μ；equal / empirical / external 权重写入协议。", "回答“对谁公平”，而不只是给一个抽象差距。"],
            ["小支持切片容易偶然成为“最差”。", "支持量、独立 cluster、tail membership 与边界质量诊断。", "避免把抽样噪声误写成公平性发现。"],
            ["不同模型可能在不同有效切片上计算。", "common-support paired comparison。", "平均性能与公平性排名反转才具有可比含义。"],
            ["普通 bootstrap 忽略空间或事件内相关。", "cluster / spatial max-T 同时风险带，并给 GeoBWER CI 与单侧 LCB。", "区分描述性差距与已认证差距。"],
            ["缺字段时容易静默退化。", "formal 模式 hard fail；exploratory 模式显式标记覆盖范围。", "提高复现性，同时保留开源可用性。"],
            ["一个数字无法说明模型整体是否好。", "联报 mean risk、tail risk、GeoBWER、支持质量与有效性状态。", "避免“所有地区都同样差，所以公平分为 0”的误读。"],
        ],
        widths=[2.05, 2.75, 1.9],
        font_size=8.15,
        header_fill=BLUE,
    )
    add_callout(
        doc,
        "风险包络解释",
        "GeoBWER 等价于允许审计者在不超过 μg/β 的约束下，把部署质量重新集中到最不利切片，"
        "再比较该最坏重加权风险与原始平均风险。这让指标具有明确的“最坏部署重加权”含义。",
        fill=PALE_CYAN,
        accent=TEAL,
        compact=True,
    )

    # Page 7 variants
    page_break(doc)
    add_title(doc, "6. 不只是一个分数：同一 GeoBWER 泛函接收不同风险输入", "Raw、Standardised、Selective、Conformal 是一个统一家族，而不是互不相干的指标")
    add_table(
        doc,
        ["模块", "它问的问题", "风险输入", "必须同时报告"],
        [
            ["Raw-GeoBWER", "哪些部署切片的实际任务风险最高？", "分类错误、BCE/Hamming、事件级分割损失等。", "平均风险、尾部风险、支持量、CI/LCB。"],
            ["Standardised-GeoBWER", "控制类别/难度组成后，地理差距还存在吗？", "按共同参考组成 π* 标准化后的切片风险。", "参考组成、缺失单元、strict/overlap/partial bounds。"],
            ["Selective-GeoBWER", "模型只处理有信心的样本后，尾部不公平是否改善？", "被接受样本上的任务损失。", "每组 coverage；零接受组必须标为不可识别。"],
            ["Conformal-GeoBWER", "达到目标覆盖率的预测集合，是否对所有切片同样可靠？", "miscoverage：真实标签不在集合中。", "总体/最差 coverage、set size、coverage debt。"],
            ["CRC-GeoBWER", "多标签或分割任务的漏标/漏检承诺，在哪些切片失效？", "每个独立单位的单调风险，如 false-negative risk。", "目标风险、违反量、效率与空间/cluster 依赖。"],
        ],
        widths=[1.25, 2.05, 1.95, 1.75],
        font_size=8.3,
        header_fill=TEAL,
    )
    add_callout(
        doc,
        "统一性的关键",
        "变的是每个样本或独立单位如何定义“风险”；不变的是“先聚合切片风险，再计算尾部风险−平均风险，并做支持与推断认证”。",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    add_formula_block(
        doc,
        ["模型输出  →  任务风险  →  部署切片风险 Rg  →  GeoBWERβ,μ  →  Audit Card"],
        fill="F7F9FC",
    )

    # Page 8 standardised/selective
    page_break(doc)
    add_title(doc, "7. Standardised 与 Selective：原有好思想如何被补强", "两者被保留，但现在明确了可识别条件和失败状态")
    add_heading(doc, "Standardised-GeoBWER：区分“组成不同”与“同类表现不同”", 2)
    add_formula_block(
        doc,
        ["Rstdg = Σc π*c · E[L | group=g, class=c]", "Standardised-GeoBWER = GeoBWERβ,μ(Rstd)"],
        fill="EAF2F8",
    )
    add_table(
        doc,
        ["以前的风险", "当前处理"],
        [
            ["某国看起来更差，可能只是因为困难类别比例更高。", "用共同参考类别组成 π* 重新加权各国风险。"],
            ["国家×类别单元缺失时，静默重归一化会改变 estimand。", "strict 模式直接标为不可识别；overlap 与 partial bounds 只作明确敏感性。"],
            ["不同模型缺失单元不同，标准化分数不可比。", "共同支持与相同参考组成下做配对比较。"],
        ],
        widths=[3.1, 3.6],
        font_size=8.8,
        header_fill=BLUE,
    )
    add_heading(doc, "Selective-GeoBWER：低风险不能以“拒绝服务”换来", 2)
    add_formula_block(
        doc,
        ["Rselg(τ) = E[L | confidence≥τ, group=g]", "coverageg(τ) = P(confidence≥τ | group=g)"],
        fill="E8F6F3",
    )
    add_body(
        doc,
        "旧版容易只看到“接受后的准确率提高”；新版强制把 coverage 一起画成风险—覆盖曲线。"
        "若某组一个样本都不接受，则该组的选择性风险不是 0，而是“不可识别”。",
    )
    add_callout(
        doc,
        "通俗例子",
        "模型在城市地区接受 90% 样本、在偏远地区只接受 5%。即使两边被接受样本同样准确，也不能称为公平：偏远地区承担了更高的服务排除。",
        fill=PALE_ORANGE,
        accent=ORANGE,
        compact=True,
    )

    # Page 9 conformal
    page_break(doc)
    add_title(doc, "8. Conformal：从“强行单选”到“诚实给出候选集合”", "它与 Selective 互补，不是简单替代关系")
    add_body(doc, "假设模型判断一张遥感图像的土地类型。普通预测只输出“草地”；Conformal Prediction 可能输出 {草地，灌木地}。")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    shape = r.add_picture(str(conformal_diagram), width=Inches(6.8))
    shape._inline.docPr.set(
        "descr",
        "从模型类别概率、独立校准集、预测集合到 GeoBWER 切片误覆盖审计的流程图。",
    )
    shape._inline.docPr.set("title", "Conformal 与 GeoBWER 审计流程")
    add_table(
        doc,
        ["方法", "模型如何表达不确定性", "公平性审计关注点"],
        [
            ["Selective", "接受一个标签，或拒绝整次预测。", "哪些切片被更频繁拒绝？接受后风险是否仍有尾部差距？"],
            ["Conformal", "输出一个可能包含多个类别的集合。", "哪些切片的真实标签更常落在集合之外？集合是否异常庞大？"],
        ],
        widths=[1.15, 2.7, 2.9],
        font_size=9.0,
        header_fill=TEAL,
    )
    add_callout(
        doc,
        "为什么不能只报 Conformal-GeoBWER？",
        "若所有地区误覆盖率都是 30%，地区差距可接近 0，但 90% coverage 承诺已经整体失败。"
        "因此必须同时报告总体 coverage、最差切片 coverage、集合大小和超过目标风险的 debt。",
        fill=PALE_RED,
        accent=RED,
    )

    # Page 10 geoconformal
    page_break(doc)
    add_title(doc, "9. 从 GeoConformal 论文到本项目：借鉴思想，但不越过理论边界", "正式理论锚点与空间局部 comparator 分层设计")
    add_heading(doc, "原论文的核心思路", 2)
    add_body(
        doc,
        "GeoConformal Prediction 将 split conformal 的校准误差按地理距离加权：离某个测试地点更近的校准样本权重更大，"
        "从而让不同地点得到不同的局部不确定性区间或预测集合。它特别适合揭示“全球平均覆盖正常、局部区域仍然失效”。",
    )
    add_heading(doc, "本项目为什么不能直接把地理核权重当作正式保证", 2)
    add_table(
        doc,
        ["理论事实", "本项目的决策"],
        [
            ["加权 conformal 在 covariate shift 下的有限样本保证，需要密度比或加权可交换性等条件。", "普通 split CP / CRC 保留为正式理论锚点。"],
            ["任意测试点中心的地理核权重，并不会自动继承上述保证；朴素局部 CP 可能欠覆盖或过覆盖。", "地理核版本准确标记为“经验空间局部 comparator”，不夸大为无条件保证。"],
            ["局部方法仍有重要诊断价值：能展示覆盖债务是否具有空间结构。", "预注册 bandwidth、只在校准集选择尺度，并报告 ESS、最近距离、可识别比例。"],
        ],
        widths=[3.2, 3.55],
        font_size=8.65,
        header_fill=BLUE,
    )
    add_callout(
        doc,
        "这不是保守退让",
        "分层命名反而增强可信度：正式 CP/CRC 回答“承诺是否有理论锚点”，地理 comparator 回答“失效是否局部化、空间化”。",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    # Page 11 task decisions
    page_break(doc)
    add_title(doc, "10. 四类任务如何实现 Conformal / GeoConformal", "统一的是审计逻辑；任务输出形式必须适配")
    add_table(
        doc,
        ["任务", "正式不确定性模块", "空间局部模块", "原因与解释"],
        [
            [
                "AlphaEarth\n单标签土地覆盖",
                "LAC / APS / RAPS prediction sets",
                "地理核 prediction-set comparator",
                "全球坐标、大样本、空间块最完整；是 GeoConformal 思想最自然的主试验台。",
            ],
            [
                "fMoW-Sentinel\n62 类分类",
                "LAC / APS / RAPS",
                "仅在坐标与局部有效样本量门控通过时运行",
                "完整 62 维概率可构造候选集合；坐标来源和局部支持必须先验证。",
            ],
            [
                "reBEN\n多标签分类",
                "Conformal Risk Control：控制漏标等单调风险",
                "做空间适用性 preflight，不强行复制单标签集合",
                "一张图本来就有多个标签；CRC 比单标签 APS 更符合任务语义。",
            ],
            [
                "Sen1Floods11\n像素分割",
                "CRC：控制事件/图块级漏检或分割风险",
                "支持和空间几何足够才进入局部比较",
                "逐像素集合容易变成大量 {洪水, 非洪水}；事件级风险控制更可解释。",
            ],
        ],
        widths=[1.2, 1.85, 1.75, 2.15],
        font_size=8.15,
        header_fill=TEAL,
    )
    add_callout(
        doc,
        "统一后的共同流程",
        "固定模型 → 使用独立 calibration 确定集合或风险阈值 → 在 test 上计算逐单位误覆盖/漏检风险 → "
        "用 GeoBWER 检查哪些部署切片承担不成比例的失败。",
        fill=PALE_BLUE,
        accent=BLUE,
    )
    add_callout(
        doc,
        "为什么不强行“四任务全覆盖同一算法”",
        "可复用不等于复制同一段公式。真正可复用的是稳定接口：任务先产生有意义的风险，再由同一 GeoBWER 协议审计。",
        fill=LIGHT_GRAY,
        accent=GRAY,
        compact=True,
    )

    # Page 12 AlphaEarth
    page_break(doc)
    add_title(doc, "11. AlphaEarth：补齐全球尺度与参考地图不确定性", "这是上次汇报后最完整的新实验轴")
    add_table(
        doc,
        ["组成", "当前设计", "回答的科学问题"],
        [
            ["数据与表征", "156,246 个样本、111 个国家；AlphaEarth embeddings；多随机种子与尺度敏感性。", "全球表征的平均性能是否掩盖国家、类别和交叉切片尾部失败？"],
            ["空间协议", "独立空间块 train / calibration / test；空间块推断。", "随机像素划分是否高估部署可靠性？"],
            ["评价参考", "ESA WorldCover 为主参考；Dynamic World 用于歧义与产品敏感性诊断。", "观察到的“错误”有多少可能来自参考地图分歧？"],
            ["公平性轴", "country、class、country×class、spatial block。", "尾部失败是地理、语义还是交互机制？"],
            ["不确定性", "split conformal + coverage/set-size GeoBWER；地理核 comparator。", "边际 coverage 是否掩盖局部 coverage debt？"],
        ],
        widths=[1.25, 2.65, 2.95],
        font_size=8.45,
        header_fill=BLUE,
    )
    add_heading(doc, "目前最值得汇报的观察", 2)
    add_bullets(
        doc,
        [
            "WorldCover 与 Dynamic World 的一致率约为 50.3%，说明土地覆盖“真值”并非无噪声；因此正式表述改为 map disagreement / evaluation-reference risk。",
            "在目标 90% coverage 下，空间块测试的边际 coverage 约为 88.99%，平均集合大小约 1.449：集合总体仍较精简，但空间转移已经造成覆盖债务。",
            "因此 AlphaEarth 不只是“第四个数据集”，而是项目中解释参考地图歧义、空间转移与审计发现能力的机制实验。",
        ],
        size=9.6,
    )
    add_callout(
        doc,
        "表述边界",
        "不能把 WorldCover 不一致全部归因于模型不公平；更严谨也更有解释力的结论是：模型—参考图分歧在不同部署切片上分布不均，并与地图产品歧义相关。",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )

    # Page 13 evidence chain/status
    page_break(doc)
    add_title(doc, "12. 四个实验现在如何组成一条论文证据链", "不是四个故事，而是四种部署切片可靠性的互补压力测试")
    add_table(
        doc,
        ["实验", "主要部署轴", "在整篇论文中的证据角色"],
        [
            ["Sen1Floods11", "灾害事件 × 传感器模态", "小 G、强相关场景：检验 fractional tail、事件 cluster 推断及强任务模型是否仍有事件尾部。"],
            ["fMoW-Sentinel", "位置/国家 × 62 类", "稀疏现实场景：检验地理泛化、平均—尾部排名反转、共同支持和模型概率质量。"],
            ["reBEN", "国家 × S1/S2/融合", "架构×模态因子场景：检验融合是否稳定改善或转移尾部风险。"],
            ["AlphaEarth", "全球国家 × 类别 × 空间块", "支持密集且全球：检验空间转移、参考图歧义、尺度敏感性与 coverage debt。"],
        ],
        widths=[1.2, 2.2, 3.5],
        font_size=8.9,
        header_fill=TEAL,
    )
    add_heading(doc, "当前科学状态：已完成什么，仍在验证什么", 2)
    add_table(
        doc,
        ["状态", "内容"],
        [
            ["已完成", "GeoBWER 公式、风险包络、协议字段、支持/无效状态、cluster/spatial 推断、common-support 比较、任务级 Conformal/CRC 代码路径。"],
            ["已有正式证据", "BWER1 四任务基础结果；AlphaEarth full v2 150k；fMoW DOFAv2 正式输出；reBEN 多条 foundation 与 supervised 路线正在完成。"],
            ["仍需冻结", "全部三随机种子、同协议 baseline、Sen1 正式面板、最终共同支持配对差值与跨任务主表。"],
            ["不应提前声称", "并非所有观察到的正 GeoBWER 都已有非零 LCB；稀疏交叉切片应称 supported-cell exploratory analysis。"],
        ],
        widths=[1.25, 5.65],
        font_size=8.9,
        header_fill=BLUE,
        first_col_fill=PALE_BLUE,
    )
    add_callout(
        doc,
        "对顶刊竞争力的实质提升",
        "提升主要来自“同一构念跨任务成立、协议公平、支持不足时不编造结论、至少一个支持密集任务能够给出认证效应”，"
        "而不是单纯把 BWER1 换成一个更复杂的数字。",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    # Page 14 next steps
    page_break(doc)
    add_title(doc, "13. 下一阶段：把方法升级转化为最终论文证据", "先完成冻结实验，再统一生成跨任务 Audit Cards")
    add_table(
        doc,
        ["优先级", "下一步", "产出"],
        [
            ["1", "完成 reBEN 剩余 supervised ResNet-50 seeds；保持已完成模型跳过与同协议比较。", "架构×模态主面板、seed 稳定性与 CRC-GeoBWER。"],
            ["2", "完成 fMoW common-9-band ResNet-50；与 DOFAv2 做共同支持配对差值。", "平均性能与尾部公平性是否排名反转的正式证据。"],
            ["3", "完成 Sen1 TerraMind / U-Net 三模态与 Prithvi 外部参考。", "事件×模态机制、CRC 与空间/事件 cluster 认证。"],
            ["4", "同步并冻结 AlphaEarth canonical lineage；补齐 GeoConformal comparator 和支持/效率报告。", "全球 coverage debt、地图歧义机制与空间局部诊断。"],
            ["5", "只用逐样本正式表做 CPU 后处理：support frontier、多分辨率轴、主表和 Audit Cards。", "主文确认性轴、附录探索轴和统一报告卡。"],
        ],
        widths=[0.65, 4.25, 2.0],
        font_size=8.6,
        header_fill=TEAL,
    )
    add_heading(doc, "最终希望形成的论文表述", 2)
    add_callout(
        doc,
        "方法贡献",
        "GeoBWER 是一个模型无关、任务感知、支持空间依赖与部署标准化的 GeoFM 审计泛函与协议，"
        "用于量化最差 β 部署质量相对平均部署风险所承担的额外风险。",
        fill=PALE_BLUE,
        accent=BLUE,
    )
    add_callout(
        doc,
        "经验贡献",
        "四任务验证显示，平均性能、尾部部署风险、选择性服务覆盖和 conformal coverage 承诺可以产生不同甚至相反的模型判断；"
        "全球平均可靠并不保证每个事件、国家、类别或传感器条件同样可靠。",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    add_callout(
        doc,
        "工具贡献",
        "最终将发布模型无关的 AuditTable、任务级 API 与 Audit Card，使第三方模型只需提供逐样本预测和部署元数据即可复用。",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )
    add_body(
        doc,
        "建议汇报收束语：这次升级没有把项目从“公平性指标”带偏，反而让最初的 BWER 从一个有直觉的实验分数，"
        "成长为能够被其他 GeoFM 论文复用、比较和质疑的正式审计框架。",
        color=BLUE,
        space_after=0,
    )

    # References
    page_break(doc)
    add_title(doc, "参考方法与汇报备注", "正文保持直观；本页用于导师追问理论来源时快速定位")
    refs = [
        (
            "Lou, Luo & Meng (2025), GeoConformal Prediction",
            "https://doi.org/10.1080/24694452.2025.2516091",
            "地理距离加权的模型无关空间预测不确定性；本项目借鉴其局部化思想。",
        ),
        (
            "Tibshirani et al. (2019), Conformal Prediction Under Covariate Shift",
            "https://proceedings.neurips.cc/paper/2019/hash/8fb21ee7a2207526da55a679f0332de2-Abstract.html",
            "说明加权 conformal 的有效性依赖加权可交换性/密度比条件。",
        ),
        (
            "Guan (2023), Localized Conformal Prediction",
            "https://doi.org/10.1093/biomet/asac040",
            "局部 conformal 需要专门构造，支持将空间局部方法与普通 split CP 分层。",
        ),
        (
            "Angelopoulos et al. (2024), Conformal Risk Control",
            "https://proceedings.iclr.cc/paper_files/paper/2024/hash/f3549ef9b5ff520a7e41ff3cc306ab2b-Abstract-Conference.html",
            "将 conformal 从预测集合覆盖推广到单调损失的期望风险控制，适合多标签与分割。",
        ),
    ]
    for title, url, note in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_hyperlink(p, title, url)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.2)
        p2.paragraph_format.space_after = Pt(7)
        r = p2.add_run(note)
        set_run_font(r, size=9.4, color=GRAY)

    add_heading(doc, "汇报使用建议", 2)
    add_bullets(
        doc,
        [
            "建议用 12–15 分钟讲完：第 1–5 节约 6 分钟，Conformal/GeoConformal 约 4 分钟，AlphaEarth 与下一步约 4 分钟。",
            "如果导师追问“为什么数值没有明显变大”，回到第 5 节：新版本的目标是更可信，不是制造更戏剧性的差距。",
            "如果导师追问“GeoConformal 是否已经完整实现”，回答：正式 split CP/CRC 已覆盖四类任务；地理核版本作为经验 comparator，优先在 AlphaEarth 运行并受支持门控。",
            "这份文档正文已经包含必要讲解，不建议再准备逐字讲稿；汇报时围绕每页绿色/橙色提示框展开即可。",
        ],
        size=9.5,
    )
    add_callout(
        doc,
        "最终边界",
        "GeoBWER 衡量的是预注册部署切片之间的尾部风险不均，不声称穷尽公平性的全部社会含义；"
        "但它能够成为 GeoFM 发布与应用时一套清晰、模型无关、可复核的公平性可靠性审计工具。",
        fill=LIGHT_GRAY,
        accent=GRAY,
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_report()
