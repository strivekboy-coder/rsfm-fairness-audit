from __future__ import annotations

import base64
import html
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def shading(element) -> str | None:
    shd = element.find(".//" + qn("w:shd"))
    if shd is None:
        return None
    fill = shd.get(qn("w:fill"))
    if not fill or fill in {"auto", "FFFFFF"}:
        return None
    return fill


def paragraph_html(paragraph, image_iter) -> str:
    p = paragraph._p
    style = paragraph.style.name if paragraph.style else "Normal"
    cls = "body"
    if style.startswith("Heading 1"):
        cls = "h1"
    elif style.startswith("Heading 2"):
        cls = "h2"
    elif "List Bullet" in style:
        cls = "bullet"
    align = paragraph.alignment
    align_css = "center" if align == 1 else "right" if align == 2 else "left"
    fill = shading(p)
    extra = f"background:#{fill};border-left:5px solid #2E75B6;padding:7px 10px;" if fill else ""
    blips = p.xpath(".//a:blip")
    if blips:
        try:
            media_path = next(image_iter)
            payload = base64.b64encode(media_path.read_bytes()).decode("ascii")
            return (
                f'<p class="image"><img src="data:image/png;base64,{payload}" '
                f'alt="embedded report diagram"/></p>'
            )
        except StopIteration:
            pass
    text = html.escape(paragraph.text).replace("\n", "<br/>")
    if not text:
        return '<p class="spacer">&nbsp;</p>'
    return f'<p class="{cls}" style="text-align:{align_css};{extra}">{text}</p>'


def table_html(table) -> str:
    rows = []
    for ri, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            fill = shading(cell._tc)
            style = f"background:#{fill};" if fill else ""
            tag = "th" if ri == 0 else "td"
            text = "<br/>".join(html.escape(p.text).replace("\n", "<br/>") for p in cell.paragraphs)
            cells.append(f"<{tag} style=\"{style}\">{text}</{tag}>")
        rows.append("<tr>" + "".join(cells) + "</tr>")
    return "<table>" + "".join(rows) + "</table>"


def main():
    source = Path(sys.argv[1]).resolve()
    output_dir = Path(sys.argv[2]).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    doc = Document(source)
    media = sorted((source.parent / "assets").glob("*.png"), key=lambda p: 0 if "bwer1" in p.name else 1)
    image_iter = iter(media)
    pages: list[list[str]] = [[]]
    para_map = {id(p._p): p for p in doc.paragraphs}
    table_map = {id(t._tbl): t for t in doc.tables}
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = para_map.get(id(child))
            if paragraph is None:
                continue
            has_page_break = any(
                br.get(qn("w:type")) == "page" for br in child.iter(qn("w:br"))
            )
            if has_page_break:
                pages.append([])
            else:
                pages[-1].append(paragraph_html(paragraph, image_iter))
        elif child.tag == qn("w:tbl"):
            table = table_map.get(id(child))
            if table is not None:
                pages[-1].append(table_html(table))

    css = """
    *{box-sizing:border-box}
    html,body{margin:0;background:#dfe5eb;font-family:"Microsoft YaHei","Segoe UI",sans-serif;color:#1F2933}
    .page{width:8.5in;height:11in;background:white;padding:.72in .82in .65in .82in;overflow:hidden;position:relative}
    .h1{font-size:16pt;color:#1F4E79;font-weight:700;margin:0 0 5pt 0;line-height:1.17}
    .h2{font-size:13pt;color:#2E75B6;font-weight:700;margin:8pt 0 4pt 0;line-height:1.15}
    .body,.bullet{font-size:10.5pt;line-height:1.10;margin:0 0 6pt 0}
    .bullet{padding-left:17pt;position:relative;margin-bottom:3pt}
    .bullet:before{content:"•";position:absolute;left:3pt;color:#2E75B6}
    .spacer{font-size:3pt;margin:0}
    table{width:100%;border-collapse:collapse;table-layout:fixed;margin:3pt 0 7pt 0;font-size:8.5pt;line-height:1.08}
    th,td{border:1px solid #cfd8e3;padding:5px 6px;vertical-align:middle;word-wrap:break-word}
    th{color:white;background:#1F4E79;font-weight:700;text-align:center}
    td:first-child{font-weight:700;color:#1F4E79}
    .image{text-align:center;margin:4pt 0 6pt 0}
    .image img{max-width:100%;max-height:3.6in;object-fit:contain}
    .footer{position:absolute;bottom:.22in;left:.82in;right:.82in;text-align:center;font-size:8pt;color:#5B6573}
    """
    for index, chunks in enumerate(pages, 1):
        document = (
            "<!doctype html><html><head><meta charset='utf-8'><style>"
            + css
            + "</style></head><body><section class='page'>"
            + "".join(chunks)
            + f"<div class='footer'>GeoBWER 项目阶段性升级汇报 · {index}</div>"
            + "</section></body></html>"
        )
        (output_dir / f"page-{index:02d}.html").write_text(document, encoding="utf-8")
    print(f"wrote {len(pages)} pages to {output_dir}")


if __name__ == "__main__":
    main()
