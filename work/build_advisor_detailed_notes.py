from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from build_advisor_geobwer_report_v2 import (
    BLUE,
    MUTED,
    OUT_DIR,
    configure_doc,
    font,
    h1,
    h2,
    h3,
    note,
    page_break,
    prose,
    table,
    title_block,
)


OUTPUT = OUT_DIR / "GeoBWER项目阶段性升级汇报_详细讲稿与答疑手册_2026-07-26.docx"


def say(doc, text: str) -> None:
    note(doc, "建议口述：", text)


def understand(doc, text: str) -> None:
    h3(doc, "自己先理解")
    prose(doc, text)


def example(doc, text: str) -> None:
    note(doc, "通俗例子：", text, fill="F5F8FB")


def caution(doc, text: str) -> None:
    note(doc, "表述边界：", text, fill="F7F7F7")


def qa(doc, question: str, answer: str) -> None:
    h3(doc, question)
    prose(doc, answer)


def source_line(doc, title: str, url: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{title}：{url}")
    font(r, size=8.8, color=MUTED)


def build() -> Path:
    doc = configure_doc("GeoBWER 阶段性升级汇报 · 详细讲稿与答疑手册")
    doc.core_properties.title = "GeoBWER项目阶段性升级汇报：详细讲稿与答疑手册"
    doc.core_properties.subject = "导师汇报逐页讲稿、任务说明、术语解释与追问答疑"
    doc.core_properties.keywords = "GeoBWER, BWER, GeoFM, RSFM, conformal prediction, CRC, AlphaEarth"

    title_block(
        doc,
        "配套讲稿 · 会前补课版",
        "GeoBWER 项目阶段性升级汇报",
        "P1–P11 逐页口述 + 四任务说明 + reBEN结果解读 + Calibration/Cluster/Conformal答疑",
    )
    prose(
        doc,
        "这份文档不是第二份汇报材料，而是给汇报者使用的“讲稿与底稿”。会议共享时仍以重制后的主报告为准；"
        "本文件用于提前理解细节、准备导师追问，并在必要时快速查找定义。正文把“可以直接讲的话”和“自己需要理解的技术细节”分开。",
    )
    note(
        doc,
        "最短使用方法：",
        "时间紧时只读每页的“建议口述”；会前重点读第 3 部分四任务档案、第 4 部分术语表和第 5–7 部分答疑。"
        "不要背公式，重点理解公式在解决什么问题。",
    )
    h2(doc, "建议汇报节奏")
    table(
        doc,
        ["内容", "建议时间", "这一段只需要让导师记住什么"],
        [
            ["P1–P2", "2–3 分钟", "项目方向未变；实验比较更受控，并新增 AlphaEarth。"],
            ["P3–P5", "4–5 分钟", "BWER1 升级为精确尾部质量的 GeoBWER；Standardised/Selective 也被统一并加上有效性门控。"],
            ["P6–P8", "4–5 分钟", "Conformal 是可靠性承诺审计；GeoConformal 是空间局部 comparator；AlphaEarth 提供全球机制证据。"],
            ["P9", "2–3 分钟", "reBEN 已完成，并首次同时展示平均、尾部、CRC 债务和选择性服务覆盖。"],
            ["P10–P11", "2–3 分钟", "四个任务是一组压力测试；reBEN 已封存，剩余面板按冻结协议完成。"],
        ],
        [1.25, 1.2, 4.05],
    )
    caution(
        doc,
        "汇报中不必讲代码版本、commit SHA 或缓存修复。只有导师追问“如何确保可复现”时，再解释代码与协议都有版本号和哈希，"
        "类似给每次正式实验加上不可混淆的身份证。",
    )

    page_break(doc)
    h1(doc, "第一部分｜P1–P11 逐页详细讲稿")

    h2(doc, "P1｜这次升级到底做了什么")
    say(
        doc,
        "上次汇报完成了 BWER1 的主体实验，核心发现是：平均性能好，不代表每个灾害事件、国家或传感器条件都可靠。"
        "这次没有偏离公平性指标方向，而是把 BWER1 重构为 GeoBWER：核心仍然是“最差尾部风险减去参考平均风险”，"
        "但现在尾部质量可以精确控制，并增加部署权重、支持量、空间相关推断、共同支持比较和置信证书。"
        "此外新增了 AlphaEarth 全球实验，以及 Conformal/CRC 不确定性审计。",
    )
    understand(
        doc,
        "这里的“公平性”不是直接判断社会群体是否受到歧视，而是研究部署切片之间的可靠性是否均衡。"
        "切片可以是国家、灾害事件、类别、季节、传感器模态或它们的交叉组合。论文的中心构念更准确地说是 "
        "deployment-slice reliability disparity：模型在最差部署尾部相对常态多承担了多少风险。",
    )
    example(
        doc,
        "一所学校平均成绩 85 分，但某个班只有 50 分。只报全校平均会掩盖问题。BWER/GeoBWER 不是重新计算全校平均，"
        "而是专门量化“最差班级群比全校常态差多少”。",
    )

    h2(doc, "P2｜为什么实验矩阵要升级")
    say(
        doc,
        "这次模型变多不是为了堆数量，而是为了让科学比较更公平。Sen1Floods11 增加同 split 的 TerraMind 与 "
        "ResNet34-U-Net 三模态比较；fMoW 用 DOFAv2 对比 common-9-band ResNet-50；reBEN 扩展为 CROMA、"
        "TerraMind 和监督 ResNet-50 的架构×模态面板。Prithvi 没有删除，而是改为任务专用强参考。"
        "AlphaEarth 则补上全球土地覆盖、空间块和参考地图歧义。",
    )
    understand(
        doc,
        "“同协议主比较”要求尽可能共享数据 split、输入波段、训练/探针预算、随机种子和审计人口。"
        "Prithvi 的 Sen1Floods11 checkpoint 已经针对该任务微调，因而很强，但不能与从相同数据起点训练的模型简单解释为纯架构优劣；"
        "所以它适合作为外部有效性强参考，而不是主因果式对比。",
    )
    table(
        doc,
        ["任务", "当前主模型/基线", "主要科学问题"],
        [
            ["Sen1Floods11", "TerraMind、ResNet34-U-Net × S1/S2/S1+S2；Prithvi TL 强参考", "模态是否同时改善平均洪水分割和事件尾部？"],
            ["fMoW-Sentinel", "DOFAv2 与 common-9-band ResNet-50，各 3 seeds", "平均准确率与地理尾部排名是否反转？"],
            ["reBEN", "CROMA、TerraMind、ResNet-50 × S1/S2/融合 × 3 seeds", "融合收益是跨架构机制，还是某个模型特例？"],
            ["AlphaEarth", "AlphaEarth 64维 embeddings + HistGradientBoosting probe", "全球地图表现、参考图歧义与空间 coverage debt 如何关联？"],
        ],
        [1.2, 2.9, 2.4],
    )

    page_break(doc)
    h2(doc, "P3｜BWER1 与 GeoBWER 的直观差别")
    say(
        doc,
        "两个版本回答同一个问题：最差部署尾部比参考平均多承担多少风险。BWER1 选择若干个完整切片；"
        "GeoBWER 则在预先定义的审计测度下，精确选择最差的 β 部署质量。"
        "例如 11 个洪水事件、β=10%：BWER1 至少要取两个完整事件，实际审计 18.18%；"
        "GeoBWER 取最差事件的全部 9.09%，再取第二差事件的 0.91%，精确组成 10%。",
    )
    understand(
        doc,
        "BWER1 的整数切片选择会随切片数量变化而跳变。GeoBWER 先给每个切片一个部署质量 μg，"
        "再从最高风险端累计到 β；边界切片可以只取一部分质量。它因此是离散切片上的 fractional upper-tail risk。"
        "当切片等权且 β=k/G 恰好是整数比例时，GeoBWER 与 BWER1 完全相同，这是嵌入性质而不是退步。",
    )
    example(
        doc,
        "把 11 个事件想成 11 个大小相同的水桶。BWER1 只能整桶搬走；要拿最差 10%，只能搬两桶。"
        "GeoBWER 可以搬走最差桶的全部，再从第二差桶取一点，刚好达到 10%。",
    )
    caution(
        doc,
        "GeoBWER 的价值不是保证每次都得到更大的数，而是让不同任务、不同切片数量和不同部署权重下的含义一致。"
        "数值与 BWER1 接近并不意味着升级无效。",
    )

    h2(doc, "P4｜新版解决了哪些旧风险")
    say(
        doc,
        "旧版最危险的不是公式会算错，而是可能把小样本偶然波动、不同模型的不同支持集、空间相关样本和稀疏交叉单元"
        "都压成一个看似精确的分数。新版增加支持门槛、共同支持配对、cluster/spatial bootstrap、同时风险带和 LCB。"
        "支持不足时会明确标为描述性、部分识别或不可识别，而不是为了出图强行给结论。",
    )
    table(
        doc,
        ["旧风险", "新版做法", "解决的含义"],
        [
            ["小组样本少却碰巧成为最差组", "支持门槛 + 同时风险带 + LCB", "区分“表面差”与“可认证差”。"],
            ["两个模型在不同群体集合上比较", "common-support paired comparison", "只在共同可比人口上比较差值。"],
            ["相邻像素/图块被当成独立样本", "事件、站点或空间块 cluster 推断", "避免置信区间过窄。"],
            ["交叉切片过稀仍被写成正式结论", "valid/descriptive/partial/invalid 状态", "诚实报告审计可发现范围。"],
        ],
        [1.8, 2.25, 2.45],
    )
    qa(
        doc,
        "LCB 是什么？",
        "LCB 是 lower confidence bound，即 GeoBWER 的置信下界。点估计为 0.20、LCB 为 0.05，表示在当前推断设计下，"
        "至少约有 0.05 的尾部额外风险得到统计认证；如果 LCB=0，只能说观察到差距，不能说已经认证差距一定大于零。",
    )
    understand(
        doc,
        "这里的 cluster 推断并不是把相邻样本合并后只预测一次。模型仍然逐张图、逐个样本或逐幅概率图完成预测，"
        "GeoBWER 点估计也仍使用全部测试样本。改变的是不确定性计算：同一事件、site、source tile 或空间块内的样本"
        "共享环境，不能假装是几百份完全独立的地理证据，因此在 bootstrap 中要作为一个相关单元共同波动。",
    )
    example(
        doc,
        "假设一个洪水事件有 100 张相邻图，另一个事件只有 5 张图。逐图 bootstrap 会把前者近似当成 100 次独立证据，"
        "区间可能虚假地很窄；cluster bootstrap 把同一事件的 100 张图一起扰动，更接近“这里只有一个事件来源”。"
        "它不会丢掉 100 张图对风险点估计的贡献，只会诚实修正我们对结论有多确定。",
    )
    qa(
        doc,
        "CI 和 LCB 到底怎么得到？",
        "正式版先按 cluster 计算各切片风险的影响，再用 multiplier bootstrap 让整个 cluster 共同获得随机权重，"
        "重复模拟切片风险的联合波动；max-T 取所有切片中最极端的标准化波动，构造同时风险带，最后把风险带传播到"
        "GeoBWER。CI 是可能范围，LCB 是该范围的下端。它不需要重新训练模型，也不会改变点预测。",
    )

    page_break(doc)
    h2(doc, "P5｜Standardised 与 Selective 为什么也要升级")
    say(
        doc,
        "Standardised-GeoBWER 用来回答：某地表现差，是因为它恰好包含更多困难类别，还是相同类别在该地仍然更差？"
        "新版要求共同可支持的类别组成，并在缺失单元时报告 overlap 或 partial bounds，不再静默重归一化。"
        "Selective-GeoBWER 则检查模型只接受高置信样本后，尾部差距是否下降；同时必须报告各组接受率，"
        "避免模型通过拒绝某些地区的大量样本制造表面公平。",
    )
    example(
        doc,
        "假设甲国样本几乎都是容易识别的机场，乙国样本很多是困难的贫民区。Raw 风险差可能主要来自任务组成。"
        "Standardised 会让两国按照同一类别配方比较。Selective 则像医院的自动分诊系统："
        "如果系统通过把乙国患者全部转给人工而获得很高准确率，必须把“乙国几乎没有人被自动服务”一起报告。",
    )
    understand(
        doc,
        "Standardisation 不是因果调整，它只能控制被测量并预先指定的组成变量。Selective prediction 的 coverage "
        "指接受率，不是 conformal coverage。前者问“模型服务了多少样本”，后者问“真实标签有多少次被预测集合覆盖”。",
    )

    h2(doc, "P6｜Conformal Prediction：从一个答案到一个候选集合")
    say(
        doc,
        "普通分类只输出一个类别；Conformal 可以输出候选集合。例如模型对土地类型不确定时输出 {草地，灌木地}。"
        "我们用独立 calibration 集决定集合应该多宽，并提前设定 90% coverage 目标。"
        "它不让模型突然变准，而是把不确定性变成可检验的覆盖承诺。"
        "GeoBWER 再检查哪些国家或类别承担了更多 miscoverage。",
    )
    example(
        doc,
        "一个模型在全球所有国家都只有 70% coverage，国家差距可能接近零，但它没有兑现 90% 承诺；"
        "另一个模型总体 90%，但某些国家只有 60%，说明承诺分配不均。因此必须同时报告总体 coverage、"
        "最差切片 coverage、集合大小和 Conformal-GeoBWER。",
    )
    caution(
        doc,
        "Conformal 的边际覆盖保证依赖 calibration 与未来样本可交换或满足相应加权条件。"
        "它不是每个国家、每个地点都自动有 90% 保证；这正是分组 coverage debt 和空间诊断存在的原因。",
    )
    understand(
        doc,
        "Calibration set 是训练结束后、正式 test 之前的一批独立有标签样本。可以把 train、calibration、test "
        "理解为学习题、模拟考试和正式考试：模型在 train 学参数；在 calibration 上只制定集合阈值或拒答阈值；"
        "test 只用于最后评价，不能反过来调阈值。BWER1 本身不需要 calibration；需要它的是 Conformal、CRC "
        "以及正式 Selective 扩展。",
    )
    example(
        doc,
        "所谓先计算“困难分数”，不是在模型预测前凭空判断困难。模型先对有真实标签的 calibration 样本输出概率，"
        "再比较真实类别的概率和排序。例如真实类别概率只有 0.08，而且排在第五名，这个样本的非符合度就较高。"
        "把全部 calibration 分数排序后取有限样本修正分位数 q，再冻结 q 到 test。",
    )
    qa(
        doc,
        "Calibration 是随机抽的吗？",
        "不是简单逐行随机抽。fMoW 按类别×site 确定性拆分，确保同一地点不跨 calibration/test；reBEN 使用官方"
        "validation；Sen1Floods11 使用官方 validation；AlphaEarth 按固定随机种子的空间块划分。共同原则是"
        "calibration 与 test 独立，并避免同一地点或空间块泄漏。",
    )

    h2(doc, "P7｜导师提到的 GeoConformal 如何融入")
    say(
        doc,
        "GeoConformal 的核心思想是按地理距离给附近校准样本更高权重，使不同测试地点拥有局部阈值。"
        "这个思想非常适合本项目，因为我们关心空间非平稳性。"
        "但任意测试点中心的地理核权重不自动继承普通 split conformal 的有限样本保证。"
        "因此当前设计保留 split CP 或 CRC 作为正式理论锚点，把地理核版本作为经验空间 comparator，"
        "比较它是否改善局部 coverage，以及是否付出更大集合或更低效率的代价。",
    )
    understand(
        doc,
        "这不是选错了论文，而是把两条相邻方法线分工：GeoConformal 提供空间局部化思想；"
        "Conformal Risk Control 提供多标签和分割任务可用的通用风险控制。"
        "单标签分类可直接构造预测集合，多标签和分割更适合控制漏检风险，而不是照搬单标签 APS。",
    )
    example(
        doc,
        "以 AlphaEarth 某个南美测试点为例：普通 split CP 用全球 calibration 样本得到同一个 q；"
        "GeoConformal 则根据地理距离，让附近 calibration 点权重更大，得到这个位置自己的 q(x)。"
        "若附近样本普遍更难，q(x) 可能更大，允许集合纳入更多候选类别；但 q 和“输出几个类别”不是一一对应，"
        "集合大小还取决于该样本整条概率分布。一个概率很集中的样本即使 q 较大，也可能只输出一个类别。",
    )
    table(
        doc,
        ["方法", "阈值如何得到", "在本项目中的角色"],
        [
            ["Split CP", "全部 calibration 分数产生一个全局 q", "单标签任务的正式边际 coverage 锚点。"],
            ["GeoConformal", "邻近 calibration 分数加权产生 q(x)", "坐标和局部支持充分时的经验空间 comparator。"],
            ["CRC", "用 calibration 选择满足目标平均风险的决策阈值", "多标签漏标与分割漏检的正式风险控制。"],
        ],
        [1.35, 2.45, 2.7],
    )
    qa(
        doc,
        "一个任务为什么会有三套预测集合？",
        "在单标签实验中，LAC、APS、RAPS 是三种不同的集合构造规则，会产生三套 test prediction sets 和三组"
        "coverage、set size、coverage debt 结果。它们共享同一模型概率和同一 calibration/test 划分，"
        "但各自独立校准阈值。多标签和分割主线通常使用 CRC 风险集合，不强行复制三套单标签算法。",
    )

    page_break(doc)
    h2(doc, "P8｜AlphaEarth 为什么重要")
    say(
        doc,
        "AlphaEarth 实验包含 156,246 个样本、111 个国家和完整类别概率，是当前最支持密集的全球实验。"
        "这里使用的不是下载并微调整个 AlphaEarth 模型，而是 Earth Engine 中由 AlphaEarth Foundations 生成的"
        " 64 维 annual Satellite Embedding 数据产品，再训练轻量 HistGradientBoosting 分类器预测 WorldCover 类别。"
        "Dynamic World 用来诊断参考地图歧义。",
    )
    understand(
        doc,
        "AlphaEarth Foundations 是产生表征的地理基础模型；Satellite Embedding V1 是该模型已经计算好的全球年度 10m、"
        "每像素 64 维数据产品；本项目抽取其中的 64 个特征，再训练下游分类器。"
        "所以“AlphaEarth embedding”既不是原始遥感影像，也不是标签，更不是我们自己训练的模型权重。",
    )
    example(
        doc,
        "可以把 AlphaEarth Foundations 想成一位已经读过海量多源遥感影像的专家；"
        "Satellite Embedding 像专家给全球每个位置写下的 64 个浓缩特征；"
        "WorldCover 是考试答案之一；我们训练的轻量分类器是拿这 64 个特征做土地覆盖答题的学生。",
    )
    caution(
        doc,
        "WorldCover 与 Dynamic World 的约 50.3% 一致率是在项目的类别映射、时间和抽样协议下得到的产品一致率，"
        "不能直接说两套地图中有一半是错的。它说明“评价参考并非无噪声真值”，需要做产品敏感性分析。",
    )
    qa(
        doc,
        "88.99% coverage 和平均集合大小 1.449 怎么解释？",
        "所有 AlphaEarth test 点仍然都被评价；88.99% 表示真实 WorldCover 类别有约 88.99% 的次数落在"
        "Conformal 候选集合内，不是只保留了 88.99% 的地点。平均集合大小 1.449 表示每个测试点平均输出约"
        "1.449 个候选类别，因此总体集合仍较精简。目标为 90% 时，总体约欠 1.01 个百分点；下一步再用"
        "Conformal-GeoBWER 检查这部分 miscoverage 是否集中在特定国家或土地类型。",
    )

    h2(doc, "P9｜reBEN 完整结果怎样讲")
    say(
        doc,
        "reBEN 的 27 条正式路线已经完成，包括三个模型、三种模态和三个随机种子。"
        "最稳定的结果是：S2 相对 S1 的可靠性优势，在三个模型、三个 seeds 和七个有支持国家组成的 63 个配对比较中全部成立。"
        "但融合并不是简单地在所有指标上都更好。例如 TerraMind 从 S2 变成 S1+S2 后，平均风险略降，尾部风险反而略升。"
        "这正好说明为什么平均性能不能代替 GeoBWER。",
    )
    understand(
        doc,
        "三 seed 均值中，CROMA 在 S1 条件最好；TerraMind 在 S2 和融合条件整体最好。"
        "九个“模型×模态”比较中有四个出现平均风险排序与 GeoBWER 排序不一致。"
        "CRC 的总体风险约 9.77%–10.70%，看起来接近 10% 目标，但最差国家仍额外承担 3.31–8.32 个百分点。"
        "Selective 的总体接受率约 50%，最差国家却只有 4.3%–24.7%。",
    )
    caution(
        doc,
        "不要说“TerraMind 的 GeoBWER 已被统计证明优于其他模型”。最保守的 GeoBWER 配对同时区间仍包含零。"
        "可以正式说的是：S1 相对 S2/S1+S2 的逐国家风险方向在全部有支持比较中一致且区间排除零；"
        "跨模型 GeoBWER 排名目前是稳定、重要但描述性的证据。",
    )

    h2(doc, "P10｜四个任务为什么不是四篇拼接研究")
    say(
        doc,
        "四个任务分别给同一框架施加不同压力：Sen1Floods11 是小样本事件与传感器模态；"
        "fMoW 是位置不相交、稀疏地理交叉和排名反转；reBEN 是大样本架构×模态机制；"
        "AlphaEarth 是全球空间转移、参考地图歧义和 coverage debt。"
        "统一点是：都先把任务输出转为逐独立单元风险，再用同一 GeoBWER 泛函审计尾部部署质量。",
    )
    understand(
        doc,
        "跨任务统一的不是原始性能指标：分割用 IoU 风险，单标签可用错误率或 log loss，多标签可用 BCE/FNR，"
        "conformal 用 miscoverage。统一的是更高一层的风险接口：每个部署单元都有可解释、同方向的损失，数值越大越差。",
    )

    h2(doc, "P11｜当前状态与下一步")
    say(
        doc,
        "GeoBWER 公式、协议与主要任务适配器已经固定，reBEN 27-run 也已经完成并可以封存。"
        "接下来顺序是完成 fMoW baseline 与 DOFA 的共同支持比较，完成 Sen1 三模态面板，"
        "再冻结 AlphaEarth 的 split CP 与空间 comparator。"
        "最后只从逐样本正式表做 support frontier、跨任务图表和可复用 Audit Cards。",
    )
    caution(
        doc,
        "可以说“方法与实现已进入冻结运行阶段”，但不要说“所有正式科学结论已经完成”。"
        "reBEN 已经有完整三 seed 证据，但整篇论文的最终结论仍依赖 fMoW、Sen1 和 AlphaEarth 的共同支持差值及置信区间。",
    )

    page_break(doc)
    h1(doc, "第二部分｜四个任务的会前补课档案")
    prose(
        doc,
        "下面的规模要分清三种口径：原始数据集规模、当前冻结实验使用规模、审计表行数。"
        "特别是 reBEN 会把一张图展开为 19 个标签决策，审计行数约为影像数的 19 倍，但影像并没有变多。",
    )

    h2(doc, "1. Sen1Floods11：全球洪水像素分割")
    table(
        doc,
        ["项目", "本项目口径"],
        [
            ["原始数据", "4,831 个 512×512、10m chips，覆盖 11 次洪水事件；其中 446 个为人工精标。"],
            ["冻结 split", "当前 TerraTorch 兼容 split 为 train/validation/test = 252/89/90，共 431；其余人工芯片不在该冻结 split 中。"],
            ["任务", "对每个像素判断洪水/水体与非洪水，属于二元语义分割，不是整张图分类。"],
            ["输入", "Sentinel-1 SAR、Sentinel-2 optical，以及 S1+S2 融合。"],
            ["模型", "TerraMind 与 ResNet34-U-Net 为同 split 主比较；Prithvi-EO-2.0 TL 为任务专用强参考。"],
            ["主要切片", "event_id（灾害事件）；模态是跨运行条件，不把 S1/S2/融合混成一个切片池。"],
            ["主要风险", "1−IoU、1−Dice、洪水漏检率；按事件聚合 TP/FP/FN/TN。"],
        ],
        [1.35, 5.15],
    )
    qa(
        doc,
        "为什么 446 与 431 不一样？",
        "446 是手工标注资产总数；252+89+90 是当前冻结官方 CSV 成员数。正式比较必须遵循同一 split，"
        "因此不能为了凑满 446 临时把未分配样本塞入 test。报告中应同时写清“数据资产总数”和“冻结评估人口”。",
    )
    qa(
        doc,
        "为什么用事件作为切片？",
        "洪水部署失败通常以事件为单位发生，同一事件内芯片共享天气、地表、成像条件和地理背景。"
        "把每个芯片当独立公平群体没有部署意义，也会忽略相关性。event_id 是操作性部署切片，不等同于受保护社会群体。",
    )

    h2(doc, "2. fMoW-Sentinel：位置不相交的 62 类场景分类")
    table(
        doc,
        ["项目", "本项目口径"],
        [
            ["冻结子集", "30,000 个样本：train 21,046；calibration 4,485；test 4,469。"],
            ["任务", "每张 Sentinel-2 影像从 62 个设施/场景类别中选一个，例如机场、港口、医院、风电场。"],
            ["split", "位置不相交；同一地点不会同时出现在训练和正式测试中。"],
            ["模型", "DOFAv2 与 common-9-band ResNet-50，各 3 seeds；保存完整 62 维概率。"],
            ["主要切片", "国家、区域、洲、纬度带、季节、类别；Country×Class 为稀疏探索轴。"],
            ["主要问题", "平均准确率排名与地理尾部 GeoBWER 排名是否分离或反转。"],
        ],
        [1.35, 5.15],
    )
    qa(
        doc,
        "为什么不直接使用原始 fMoW-Sentinel 全量几十万图像？",
        "本项目使用的是预先构建的 30k 支持感知、位置不相交审计子集，目的是保证地理元数据、位置隔离与跨模型可复现。"
        "它不是宣称重建完整 fMoW benchmark，而是构建适合公平性审计的受控人口。",
    )
    qa(
        doc,
        "为什么 Country×Class 有时只剩很少有效单元？",
        "62 类乘以许多国家后，理论组合数量很大，但很多国家从未出现某些设施，或只有极少独立地点。"
        "“国家×类别”不是数学乘法，而是交叉单元，例如“乌干达×机场”。"
        "数据量总体很大也不保证每个交叉格子有足够支持。",
    )

    page_break(doc)
    h2(doc, "3. reBEN / BigEarthNet v2：大规模多标签土地覆盖")
    table(
        doc,
        ["项目", "本项目口径"],
        [
            ["原始数据", "549,488 对配准的 Sentinel-1 / Sentinel-2 patches，来自欧洲 10 个国家。"],
            ["任务", "19 类多标签场景分类；一张图可同时包含农田、森林、草地、水体等多个标签。"],
            ["标签来源", "CORINE Land Cover 2018 派生的 19 类体系，并带改进的地理 split。"],
            ["模型矩阵", "CROMA、TerraMind、supervised ResNet-50 × S1/S2/S1+S2 × seeds 42/73/101，共 27 条正式路线。"],
            ["主要切片", "国家、类别、国家×类别；传感器模态是实验条件。"],
            ["审计行数", "每张图对 19 个标签各产生一个决策行；因此 sample×label 行数不是影像数。"],
            ["主要问题", "融合是否跨架构降低平均风险与国家尾部，还是只把风险转移到不同类别/国家。"],
        ],
        [1.35, 5.15],
    )
    qa(
        doc,
        "多标签与单标签到底差在哪？",
        "fMoW 一张图只有一个正确类别；reBEN 一张图可以同时有多个正确标签。"
        "因此 fMoW 的 prediction set 问“真实唯一类别是否在候选集合中”；reBEN 的 CRC 问“真实标签中平均漏掉了多少”。",
    )
    qa(
        doc,
        "CROMA、TerraMind、ResNet-50 分别扮演什么角色？",
        "CROMA 是专门学习雷达—光学对齐表征的多模态基础模型；TerraMind 是更通用的多模态生成式 EO 基础模型；"
        "ResNet-50 是从监督数据训练的常规深度学习基线。三者同时出现，才能判断融合收益是基础模型共性、"
        "某个预训练架构特性，还是普通监督模型也能获得的收益。",
    )
    qa(
        doc,
        "119,825 个测试样本很多，为什么区间仍然可能很宽？",
        "因为空间推断不能把每张相邻影像当成完全独立的地理证据。reBEN 的正式推断依赖 45 个 source-tile clusters；"
        "同一 tile 内的大量样本共享地理环境。119,825 决定点估计很稳定，但 45 个独立 cluster 更直接决定地理差异区间有多窄。",
    )
    qa(
        doc,
        "为什么只报告七个国家？97.7% 样本覆盖不是已经接近全部吗？",
        "固定全集有十国，但 Kosovo、Luxembourg、Switzerland 没达到主要独立支持门槛。七国包含 97.7% 的样本，"
        "却只占“每个国家等权”部署测度的 70%。这说明样本覆盖与部署质量覆盖不是一回事。"
        "正式点估计只能称为七国支持全集 GeoBWER；完整十国只能给很宽的 partial bounds。",
    )
    qa(
        doc,
        "为什么 63/63 个国家比较显著，模型 GeoBWER 差异却可能不显著？",
        "逐国家差值是在固定国家上比较两个风险；GeoBWER 差值还要先识别最差尾部、传播组风险不确定性，并同时控制多组比较。"
        "后者是更困难、更保守的问题。因此可以有非常稳定的模态方向性，同时仍不足以严格排序两个模型的 GeoBWER。",
    )
    qa(
        doc,
        "“平均排名与 GeoBWER 排名 4/9 不一致”说明什么？",
        "说明平均风险与地理尾部不均衡相关，但不是同一个信息。相关性约为 0.875，并非毫无关系；"
        "真正价值在于某些模型平均略好时，尾部差距可能反而更大。它支持 GeoBWER 是补充指标，而不是替代总体性能。",
    )
    qa(
        doc,
        "reBEN 的 CRC 和 Selective 结果为什么很重要？",
        "CRC 显示总体风险接近 10% 目标时，最差国家仍额外承担约 3–8 个百分点，证明平均风险控制不等于地理公平。"
        "Selective 显示总体接受约一半样本时，最差国家可能只接受 4.3%–24.7%，证明低风险也可能来自不均衡拒答。"
        "二者把 GeoBWER 从普通性能差距扩展到“可靠性承诺”和“服务可得性”差距。",
    )
    qa(
        doc,
        "reBEN 已经做了导师所说的 GeoConformal 吗？",
        "完成的是多标签 CRC 与国家切片尾部审计。连续坐标的地理核 localization 因 calibration/test 缺少可靠坐标而 27/27 均被 preflight 标记为 screened_not_run。"
        "这不影响 CRC 主结论，但汇报时不能说 reBEN 已完成连续空间 GeoConformal；真正适合该扩展的是坐标完整的 AlphaEarth。",
    )

    h2(doc, "4. AlphaEarth：全球表征产品上的土地覆盖审计")
    table(
        doc,
        ["项目", "本项目口径"],
        [
            ["冻结样本", "156,246 个样本，111 个国家，带坐标、空间块、完整类别概率。"],
            ["表征", "Google Satellite Embedding V1：每个 10m 像素为 64 维、年度、多源 EO 表征。"],
            ["模型关系", "表征由 AlphaEarth Foundations 生成；本项目训练 HistGradientBoosting 下游分类器，不训练 AlphaEarth 权重。"],
            ["主参考", "ESA WorldCover：全球 10m、11 类年度土地覆盖图。"],
            ["歧义参考", "Dynamic World：Sentinel-2 驱动的近实时 10m、9 类概率与标签产品。"],
            ["主要切片", "国家、土地覆盖类别、国家×类别、空间块，以及产品一致/分歧区域。"],
            ["主要问题", "全球平均性能是否掩盖国家/类别尾部；coverage debt 是否空间集中；参考产品分歧解释多少表面错误。"],
        ],
        [1.35, 5.15],
    )
    qa(
        doc,
        "AlphaEarth embedding 到底是数据集还是模型？",
        "严格说有两个对象：AlphaEarth Foundations 是模型；Satellite Embedding V1 是该模型生成并发布在 Earth Engine 的数据产品。"
        "本项目直接使用后者，所以代码输入看起来像一个 64 波段数据集；但其科学意义来自前者学到的表征。",
    )
    qa(
        doc,
        "为什么还要训练一个 HistGradientBoosting？",
        "Embedding 只是一串可迁移特征，不会自动输出 WorldCover 类别。轻量分类器把 64 维表征映射到具体任务标签。"
        "这相当于冻结基础模型并训练下游 probe，可用于测量表征在不同国家和类别上的可用性。",
    )

    page_break(doc)
    h1(doc, "第三部分｜模型、数据与实验单位术语表")

    h2(doc, "模型、数据集、表征、标签：四者不要混")
    table(
        doc,
        ["对象", "含义", "本项目例子"],
        [
            ["基础模型", "在大规模 EO 数据上预训练、可迁移到下游任务的参数化模型。", "DOFAv2、CROMA、TerraMind、Prithvi、AlphaEarth Foundations"],
            ["数据集", "影像、标签和元数据的样本集合。", "Sen1Floods11、fMoW-Sentinel、reBEN"],
            ["表征/embedding", "模型把原始输入压缩成的特征向量。", "DOFA 768维 embedding；AlphaEarth 64维 annual embedding"],
            ["下游模型/probe", "利用 embedding 完成具体任务的轻量模型。", "线性 probe、HistGradientBoosting"],
            ["参考标签/地图", "用于评价预测的目标或代理真值。", "洪水手工 mask、fMoW 类别、CLC 标签、WorldCover"],
        ],
        [1.35, 2.25, 2.9],
    )

    h2(doc, "当前主要模型的通俗介绍")
    h3(doc, "DOFAv2")
    prose(
        doc,
        "DOFA 的核心是根据波长/通道信息动态适配不同遥感传感器，而不是只接受固定 RGB。"
        "在 fMoW 中使用固定 9 个共同 Sentinel-2 波段，使其与 ResNet-50 的输入尽量对齐。"
    )
    h3(doc, "CROMA")
    prose(
        doc,
        "CROMA 同时学习 Sentinel-1 雷达与 Sentinel-2 光学表征，并有专门的融合编码器。"
        "它非常适合 reBEN 的 S1、S2、S1+S2 受控模态问题。"
    )
    h3(doc, "TerraMind")
    prose(
        doc,
        "TerraMind 是面向多种 EO 模态的 any-to-any 生成式基础模型。"
        "本项目利用其统一多模态骨干，在 Sen1Floods11 和 reBEN 上比较 SAR、光学与融合输入。"
    )
    h3(doc, "Prithvi-EO-2.0")
    prose(
        doc,
        "Prithvi 是 IBM/NASA 的多时相 EO 基础模型。Sen1Floods11 路线使用官方任务适配后的分割 checkpoint，"
        "因此它是强而有现实意义的外部参考，但其训练历史与同 split 从头训练的基线不完全相同。"
    )
    h3(doc, "ResNet-50 与 ResNet34-U-Net")
    prose(
        doc,
        "它们不是基础模型，而是监督学习基线。ResNet-50 用于场景分类；U-Net 解码结构用于逐像素分割。"
        "基线不是“低级模型”，而是检验预训练基础模型是否真的在尾部可靠性上带来额外价值的必要参照。"
    )

    h2(doc, "切片、切片×切片、实验条件")
    qa(
        doc,
        "什么是切片（slice）？",
        "切片是预先定义、具有部署意义的一组样本，例如某个国家、某次洪水事件、某一季节或某种类别。"
        "切片不是算法自动随意找出的最差十几条样本，而是可解释、可复核的部署单元。",
    )
    qa(
        doc,
        "切片×切片是什么意思？",
        "表示两个轴的交叉单元，不是数值相乘。例如 Country×Class 中，“印度×机场”和“印度×医院”是两个不同格子。"
        "交叉切片能发现边际平均掩盖的问题，但维度越高越稀疏，所以必须报告有效单元比例与样本覆盖。",
    )
    qa(
        doc,
        "S1、S2、S1+S2 为什么通常不是一个切片轴？",
        "它们在 reBEN/Sen1 中代表模型输入条件和独立运行路线。正确比较是：在每种模态内部计算国家/事件 GeoBWER，"
        "再比较三种条件；如果把三种模态混进一个群体池，模型和数据条件会混在一起。",
    )
    qa(
        doc,
        "独立单元、cluster、切片有什么区别？",
        "独立单元是最小不重复评价对象，例如 site、chip 或 event；cluster 是推断时共同重采样的相关单元，"
        "例如同一地点或空间块；切片是要比较的部署群体。一个国家切片可以包含很多 site clusters。",
    )

    page_break(doc)
    h1(doc, "第四部分｜GeoBWER 家族：从一个分数到一套审计协议")

    h2(doc, "GeoBWER 的最小数学理解")
    prose(
        doc,
        "对每个切片 g 先计算风险 Rg，并给它部署质量 μg。把切片按风险从高到低排序，"
        "取累计质量恰好为 β 的最差尾部，得到 Tailβ,μ(R)。再减去同一测度下的参考平均 mμ(R)："
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("GeoBWERβ,μ = Tailβ,μ(R) − mμ(R)")
    font(r, size=14, bold=True, color=BLUE)
    prose(
        doc,
        "β 表示关注最差多少部署质量，例如 10%；μ 表示“部署世界如何加权”，可以是切片等权、"
        "按观测人口加权，或由外部目标部署分布给权。风险越大越差，因此 GeoBWER 越大，尾部相对平均越糟。",
    )
    qa(
        doc,
        "为什么不是只报最差组？",
        "单个最差组对噪声、切片细分和小样本很敏感。尾部平均保留了最坏部署关注，同时比单一最大值稳定。"
        "β-profile 还可以从极端最差逐步看到更宽尾部。",
    )
    qa(
        doc,
        "为什么 GeoBWER=0 仍可能是坏模型？",
        "因为它度量差距，不度量绝对合格性。如果所有国家错误率都是 40%，它们同样差，GeoBWER 可接近 0。"
        "所以 Audit Card 必须同时报告总体平均风险、尾部绝对风险、GeoBWER 和目标违反量。",
    )
    qa(
        doc,
        "equal、empirical、external 权重有什么区别？",
        "equal 表示每个切片在审计中同等重要；empirical 表示按当前数据中的样本/部署质量加权；"
        "external 表示按预先给定的真实部署人口加权。三者回答不同问题，不能混用后比较。",
    )

    h2(doc, "Raw、Standardised、Selective、Conformal/CRC 的统一关系")
    table(
        doc,
        ["版本", "送入 GeoBWER 的风险", "回答的问题"],
        [
            ["Raw-GeoBWER", "原始任务损失", "哪些部署切片的实际风险最高？"],
            ["Standardised-GeoBWER", "按共同类别/组成分布标准化后的组风险", "控制任务组成后，地理差距是否仍存在？"],
            ["Selective-GeoBWER", "只在被模型接受的样本上的任务风险", "拒答后，保留服务中的尾部差距是否下降？"],
            ["Conformal-GeoBWER", "预测集合是否漏掉真实标签（miscoverage）", "覆盖承诺是否在不同切片公平兑现？"],
            ["CRC-GeoBWER", "多标签/分割的每单元漏检风险", "应用风险控制是否在不同群体发生尾部违反？"],
        ],
        [1.55, 2.25, 2.7],
    )
    qa(
        doc,
        "Standardised-GeoBWER 是不是被删除了？",
        "没有。它不再是与 GeoBWER 平行的另一个核心公式，而是先生成标准化风险 Rstd,g，再交给同一个 GeoBWER 泛函。"
        "这样 Raw、Standardised、Selective、Conformal 的差别来自风险输入，尾部定义保持一致。",
    )
    qa(
        doc,
        "strict standardisation、overlap 和 partial bounds 是什么？",
        "strict 要求预注册组成单元都有足够支持；overlap 只在各组共同出现的组成上比较；"
        "partial bounds 对缺失单元给出最坏/最好可能区间。它们分别强调完整目标、可比子人口和不完全识别。",
    )

    page_break(doc)
    h1(doc, "第五部分｜LAC、APS、RAPS 与 CRC 的通俗解释")

    h2(doc, "先分清三个词：coverage、set size、efficiency")
    prose(
        doc,
        "Coverage 是真实标签落在预测集合中的比例；set size 是每次输出多少个候选；efficiency 通常指在满足 coverage 的同时集合尽量小。"
        "一个把 62 类全部输出的系统 coverage 很高，但没有决策价值，所以不能只看 coverage。",
    )
    example(
        doc,
        "目标 coverage=90%。模型 A 平均输出 1.4 个类别并达到 90%；模型 B 每次输出 62 类也达到 100%。"
        "B 的 coverage 更高，但不确定性表达毫无效率。",
    )

    h2(doc, "LAC：按每个类别自己的概率门槛入选")
    prose(
        doc,
        "LAC 在当前实现中使用真实类别非符合度 s=1−p(y)。校准集给出阈值 q；测试时把所有满足 1−p(k)≤q 的类别放入集合。"
        "直观上就是“类别概率足够高就入选”。它简单、集合通常较小，但对困难样本的自适应性有限。",
    )
    example(
        doc,
        "若校准后要求类别概率至少 0.20，测试概率为草地 0.55、灌木 0.30、农田 0.10，"
        "LAC 输出 {草地，灌木}。",
    )

    h2(doc, "APS：按累计概率质量构造自适应集合")
    prose(
        doc,
        "APS 先把类别按概率从高到低排序，再逐个加入高概率类别，直到校准的累计质量阈值。"
        "简单样本可能只输出一个类别，困难样本会自动输出更多类别，因此比固定概率门槛更能适应样本难度。",
    )
    example(
        doc,
        "概率为 0.55、0.30、0.10、0.05。若校准阈值要求覆盖约 0.80 的前序概率质量，"
        "通常会包含前两个或边界类别；具体边界遵循项目冻结的 deterministic APS 规则。",
    )

    h2(doc, "RAPS：给 APS 的长尾集合加正则")
    prose(
        doc,
        "RAPS 是 Regularized Adaptive Prediction Sets。它在 APS 的累计概率分数上，对排名超过 kreg 的低位类别"
        "加入 λ 惩罚，防止困难样本把大量极低概率类别都塞进集合。它在 coverage 与集合大小之间更实用。",
    )
    note(
        doc,
        "术语纠正：",
        "项目代码和正式协议中使用的是 RAPS，不是 PAPS。如果汇报时看到“PAPS”，应视为口误或拼写错误，"
        "不要把它讲成第四种已实现方法。",
    )
    table(
        doc,
        ["方法", "最通俗规则", "优点", "代价/注意"],
        [
            ["LAC", "每个类别概率够高就入选", "简单、易解释、集合常较小", "对样本难度自适应较弱"],
            ["APS", "从最高概率开始累计", "困难样本自动扩大集合", "类别多时集合可能变大"],
            ["RAPS", "APS + 低排名类别惩罚", "抑制巨大集合、效率更好", "多出 λ 与 kreg，但必须预注册"],
        ],
        [0.8, 2.0, 1.8, 1.9],
    )

    h2(doc, "多标签 CRC：控制一张图漏掉多少真实标签")
    prose(
        doc,
        "reBEN 的真实答案本来就是多个标签。项目在 calibration 集上寻找一个概率阈值，使有限样本修正后的平均 false-negative risk "
        "不超过 α。测试时每个标签概率高于阈值就预测为存在，并计算每张图漏掉的真实标签比例。",
    )
    example(
        doc,
        "真实标签是 {农田，草地，水体}，模型输出 {农田，水体}，则漏掉 1/3，false-negative risk=0.333。"
        "CRC 的目标不是保证每张图都低于 0.10，而是控制未来可交换样本上的期望风险；GeoBWER 再检查哪些国家的风险特别高。",
    )

    h2(doc, "分割 CRC：控制每张图/事件的洪水漏检比例")
    prose(
        doc,
        "Sen1Floods11 中每张图有成千上万像素。项目不把像素当作相互独立的 conformal 样本，"
        "而是用 calibration 概率图选择像素阈值，计算每张图真实洪水像素中有多少被漏掉，并在事件层面审计尾部。"
    )
    example(
        doc,
        "一张图有 1,000 个真实洪水像素，预测漏掉 80 个，FNR=8%。另一事件平均漏掉 35%，"
        "即使全局平均达标，CRC-GeoBWER 仍会揭示后一个事件承担的 coverage/risk debt。",
    )
    caution(
        doc,
        "CRC 控制的是预先定义的单调风险。若论文关心误报、IoU 或成本加权损失，需要另行冻结对应风险函数；"
        "不能看到结果后再换成最有利的风险。",
    )

    page_break(doc)
    h1(doc, "第六部分｜跨页串联复习：从 Calibration 到 GeoBWER 的完整计算链")
    prose(
        doc,
        "这一部分不对应汇报中的某一张单页，而是把 P4 的空间推断、P6 的 calibration、P7 的 GeoConformal/CRC "
        "和 P8 的 AlphaEarth 结果串成一条完整流程。逐页讲解时先看前面的对应 P 页；需要深入理解或回答追问时再查本节。",
    )

    h2(doc, "Calibration 样本到底是什么")
    prose(
        doc,
        "Calibration set 是模型训练完成后、正式 test 之前的一批独立有标签样本。它不参与模型参数训练，也不能与 test 重叠；"
        "它的作用是提前冻结 prediction set、CRC 或 Selective 所需的部署阈值。"
        "因此它不是模型跑完后专门挑出的困难样本，而是在查看 test 结果之前已经由 split 协议确定。",
    )
    example(
        doc,
        "可以把 train、calibration、test 想成“学习题、模拟考试和正式考试”。Train 用来学习；calibration 用来决定正式考试中"
        "“不确定到什么程度可以多报一个答案”；test 只用于最后检验，不能偷看 test 答案后再调整规则。",
    )
    table(
        doc,
        ["任务", "Calibration 来源", "为什么不是普通逐样本随机抽取"],
        [
            ["fMoW-Sentinel", "从原 holdout 按类别×地点拆出 calibration；固定 seed", "同一 site 的多时相影像整体进入同一 split，calibration 与 test 地点不重叠。"],
            ["reBEN", "官方 validation split", "官方 test 保持封闭；监督模型选择使用 train 内部 group-disjoint holdout。"],
            ["Sen1Floods11", "官方 validation：89 个芯片", "与正式 test 的 90 个芯片分离，按冻结成员关系使用。"],
            ["AlphaEarth", "固定 seed 的空间块级 calibration", "同一空间块不能跨 calibration/test，避免相邻位置泄漏。"],
        ],
        [1.25, 2.25, 3.0],
        size=8.8,
    )
    qa(
        doc,
        "BWER1 为什么没有 calibration？",
        "BWER1 的核心只读取 test 损失与切片标签，例如国家错误率或事件 1−IoU，不需要选择 prediction-set 阈值。"
        "Calibration 不是 GeoBWER 数学本身的要求，而是新版 Conformal、CRC 和正式 Selective 扩展的要求。",
    )

    h2(doc, "“困难分数”为什么能在 calibration 上计算")
    prose(
        doc,
        "这里的“第一步计算困难分数”是指模型已经完成 calibration 推理之后，Conformal 后处理的第一步。"
        "Calibration 同时拥有模型概率和真实标签，因此可以观察真实类别被模型放在多高的位置。"
        "真实类别概率高、排名靠前，非符合度小；真实类别概率低、排名靠后，非符合度大。",
    )
    example(
        doc,
        "两张 fMoW calibration 图都被 top-1 预测错。图 A 的真实“机场”概率为 0.48、只比第一名 0.51 稍低；"
        "图 B 的真实“机场”概率只有 0.01。Accuracy 都记作一次错误，但 Conformal 会认为图 B 明显更困难。",
    )
    table(
        doc,
        ["方法", "Calibration 中的困难/非符合度定义", "最终阈值"],
        [
            ["LAC", "真实类别分数 1−p(y)", "全局 qLAC 或地理局部 qLAC(x)"],
            ["APS", "排在真实类别之前的累计概率质量", "全局 qAPS 或地理局部 qAPS(x)"],
            ["RAPS", "APS 分数 + 对过低排名类别的正则惩罚", "全局 qRAPS 或地理局部 qRAPS(x)"],
            ["CRC", "某个概率阈值下的平均漏标/漏检风险", "满足有限样本风险目标的概率阈值 t"],
        ],
        [0.9, 3.6, 2.0],
        size=8.8,
    )

    page_break(doc)
    h2(doc, "q 如何决定一个样本输出几个类别")
    prose(
        doc,
        "q 不是“输出类别数”，而是允许的最大非符合度。测试样本的 prediction set 为"
        " C(x)={y:s(x,y)≤q}。q 越大，通常有更多候选类别通过；但最终进入几个仍取决于该样本自己的完整概率分布。",
    )
    example(
        doc,
        "假设五个候选类别的非符合度依次为 0.00、0.43、0.71、0.87、0.95。"
        "q=0.50 时输出前两个；q=0.75 时输出前三个；q=0.90 时输出前四个。"
        "另一张图的分数排列不同，同一个 q 可能只输出一个，也可能输出五个。",
    )
    qa(
        doc,
        "一个任务运行 LAC、APS、RAPS，是否会产生三套输出？",
        "是。单标签任务会分别得到 CLAC(x)、CAPS(x)、CRAPS(x)，并各自报告 coverage、平均 set size、miscoverage GeoBWER 和 efficiency GeoBWER。"
        "如果启用地理核 comparator，还会对应产生 Geo-LAC、Geo-APS、Geo-RAPS。reBEN 与 Sen1 使用 CRC，不使用单标签三套集合。",
    )

    h2(doc, "Split CP、GeoConformal 与 CRC 的关系")
    table(
        doc,
        ["方法", "阈值怎样产生", "用于什么任务", "在项目中的地位"],
        [
            ["Split CP", "所有 calibration 分数共同产生一个全局 q", "AlphaEarth、fMoW 单标签", "正式边际 coverage 理论锚点"],
            ["GeoConformal", "按测试地点距离对 calibration 分数加权，产生 q(x)", "坐标与局部支持充分的单标签任务", "经验空间局部 comparator"],
            ["CRC", "选择使有限样本修正平均风险≤α的概率阈值 t", "reBEN 多标签、Sen1 分割", "任务正确的正式风险控制锚点"],
        ],
        [1.05, 2.75, 1.65, 1.55],
        size=8.7,
    )
    prose(
        doc,
        "GeoConformal 不是 CRC 在本项目中的体现。二者共享“train→calibration→冻结阈值→test”范式，但控制对象不同。"
        "GeoConformal 主要局部化单标签 prediction set；CRC 控制一般的平均漏标或漏检风险。"
        "在 reBEN/Sen1 中，连续坐标局部 CRC 尚没有被当作具有正式保证的方法，因此保留全局 CRC，再由 GeoBWER 审计国家或事件尾部。",
    )
    note(
        doc,
        "最简流程图：",
        "模型先输出概率；任务适配器用 calibration 冻结 LAC/APS/RAPS 或 CRC 阈值；在 test 上产生 miscoverage/FNR；"
        "GeoBWER 最后审计这些风险在国家、事件或类别之间如何分配。GeoBWER 本身不训练模型，也不选择 q。",
    )

    h2(doc, "Calibration 是否贯穿整个 GeoBWER")
    table(
        doc,
        ["分析", "是否需要 conformal calibration", "原因"],
        [
            ["Raw-GeoBWER", "否", "直接审计冻结 test 任务损失。"],
            ["Standardised-GeoBWER", "否", "对 test 组风险按共同组成重新加权。"],
            ["β-profile / support / common support", "否", "属于 GeoBWER 估计对象与可比人口。"],
            ["CI / LCB", "否", "由 test 审计表与独立 cluster 的抽样波动计算。"],
            ["Selective-GeoBWER", "当前正式实现需要", "用 calibration 冻结目标接受率对应的置信度阈值。"],
            ["Split/Geo-Conformal", "需要", "用 calibration 决定 q 或 q(x)。"],
            ["CRC-GeoBWER", "需要", "用 calibration 决定满足目标风险的概率阈值。"],
        ],
        [2.1, 1.8, 3.1],
        size=8.7,
    )
    prose(
        doc,
        "很多新版扩展需要 calibration，是因为它们都在制定一个部署决策规则；凡是需要从数据中选择阈值，又不能使用 test 标签，"
        "就需要独立 calibration。GeoBWER 核心仍然只是一个风险聚合与尾部审计泛函。",
    )

    page_break(doc)
    h2(doc, "AlphaEarth 的 88.99% coverage 与 1.449 set size")
    prose(
        doc,
        "Coverage 不是“保留了多少地点”，而是所有测试点都得到 prediction set 后，真实 WorldCover 类别被包含在集合中的比例。"
        "88.99% 表示约每 1,000 个测试点中有 890 个集合包含真实类别，约 110 个没有包含。",
    )
    prose(
        doc,
        "平均 set size=1.449 表示每个测试点平均输出 1.449 个候选类别。单个集合大小仍是 1、2、3 等整数；"
        "1.449 是跨样本平均，说明集合总体没有为了提高 coverage 而无限膨胀。",
    )
    qa(
        doc,
        "“欠覆盖”是什么？",
        "目标 coverage=90%，实际=88.99%，所以总体 coverage debt 为 1.01 个百分点；等价地，允许 miscoverage 为 10%，实际为 11.01%。"
        "边际 coverage 只是全体平均，某国若只有 70%，该国 debt 为 20 个百分点。这正是 Conformal-GeoBWER 要寻找的风险集中。",
    )

    h2(doc, "Cluster 推断：不是把相邻图片合并后只预测一次")
    prose(
        doc,
        "模型仍然对每张图片或芯片预测，点估计也使用全部样本。变化只发生在统计不确定性估计："
        "同一事件、site、source tile 或空间块内的样本共享环境与成像条件，不能被当作许多完全独立的地理重复。",
    )
    example(
        doc,
        "某国有 6 张图：1–3 来自 tile A，4–5 来自 tile B，6 来自 tile C。国家平均风险仍用 6 张图；"
        "但 CI 主要依据 A、B、C 三个相对独立来源之间的波动，而不是假装拥有 6 份独立地理证据。",
    )
    qa(
        doc,
        "为什么这样能解决“相邻样本虚假增加证据量”？",
        "如果 tile A 很容易，它产生的三张相似图不应被当成三次独立成功验证。Cluster 推断让 tile A 的样本在每次扰动中共同变化，"
        "因此同一来源不能凭图片数量重复投票。相关样本集中在少数 clusters 时，CI 会更宽，也更符合真实证据量。",
    )
    table(
        doc,
        ["任务", "主要推断 cluster", "切片示例"],
        [
            ["reBEN", "source tile / MGRS tile", "国家"],
            ["fMoW-Sentinel", "category-scoped site", "国家、区域、类别"],
            ["Sen1Floods11", "场景/空间块", "洪水事件"],
            ["AlphaEarth", "空间块", "国家、土地覆盖类别"],
        ],
        [1.4, 2.4, 2.7],
    )

    h2(doc, "当前正式 bootstrap、max-T、CI 与 LCB 怎么连接")
    prose(
        doc,
        "当前正式 GeoBWER 1.1 主要使用 cluster/spatial multiplier bootstrap，而不是重新训练模型 2,000 次，也不是简单逐行有放回抽样。"
        "程序先计算每个 cluster 对各切片风险的影响，再给整个 cluster 随机赋予 +1/−1 multiplier，重复约 2,000 次，模拟所有切片风险会怎样共同波动。",
    )
    prose(
        doc,
        "每次重复都记录所有预注册切片中最大的标准化波动；其 95% 分位数形成 max-T 临界值。"
        "因此得到的是所有国家/事件同时成立的风险带，而不是分别为每个国家做一个未经多重比较控制的窄区间。",
    )
    prose(
        doc,
        "随后把同时风险带通过 GeoBWER 的 Lipschitz/风险包络界传播，得到 GeoBWER 的保守 CI。"
        "LCB 是该区间下端点：LCB>0 表示至少有一部分尾部超额风险得到认证；LCB=0 只表示当前独立 cluster 支持不足以排除零，"
        "不能解释为模型已经公平或两个模型等价。",
    )
    note(
        doc,
        "一句话区分两个“空间”模块：",
        "GeoConformal 利用邻近 calibration 样本改变每个地点的 prediction set；cluster/spatial max-T 利用相关单元修正 GeoBWER 的 CI/LCB。"
        "前者管预测集合，后者管结论可信度，两者不是同一个算法。",
    )

    page_break(doc)
    h1(doc, "第七部分｜WorldCover、Dynamic World 与参考地图歧义")

    h2(doc, "两套产品分别是什么")
    table(
        doc,
        ["产品", "性质", "类别/时间", "在项目中的角色"],
        [
            ["ESA WorldCover", "Sentinel-1+2 制作的全球 10m 离散土地覆盖图", "2020/2021；11 类", "主要评价参考标签"],
            ["Dynamic World", "逐幅 Sentinel-2 生成的近实时 10m 概率与 top-1 标签", "2015年至今；9 类", "歧义、置信度和产品敏感性诊断"],
        ],
        [1.25, 2.25, 1.35, 1.65],
    )
    prose(
        doc,
        "两套产品的传感器、时间聚合、类别体系和算法都不同。WorldCover 更像年度离散地图；"
        "Dynamic World 更像随 Sentinel-2 观测变化的概率序列。因此必须先做类别 crosswalk，再比较一致/分歧。",
    )
    qa(
        doc,
        "为什么这件事对公平性重要？",
        "如果某些国家更多是草地—灌木地过渡区、季节性农田或混合像素，参考产品更容易分歧。"
        "把所有“模型≠WorldCover”都叫模型偏见，会把标签构念误差错误归因给模型。"
        "审计应区分 model-reference disagreement 与 reference-product ambiguity。",
    )
    qa(
        doc,
        "约 50.3% 一致率是不是说明地图很差？",
        "不能这样说。这是经过项目 taxonomy 映射、抽样年份与空间协议后的产品一致率，不是两张图的官方精度对比。"
        "它可能同时包含类别定义差异、时间差、混合像素、真实变化和产品误差。正确结论是参考不确定性不可忽略。",
    )
    qa(
        doc,
        "只有“用 Dynamic World 做歧义诊断”这一种处理吗？",
        "不是。可选方案还包括：只在两产品一致区做确认性分析；用 Dynamic World 概率作为软标签；"
        "建立潜在真值/标签噪声模型；抽样做人类专家复核；加入第三套地图做多参考敏感性；"
        "对不确定区域报告区间而非硬错误。本项目选择“WorldCover 主参考 + Dynamic World 分层敏感性”是因为它全球可扩展、"
        "模型无关并保留明确主 estimand；其他方法适合作为机制或附录扩展。",
    )
    example(
        doc,
        "WorldCover 标为草地，Dynamic World 高概率标为灌木，模型也预测灌木。旧分析记作模型错误；"
        "新版先标记为产品分歧区，再问：模型尾部差距在两产品一致区是否仍然存在。"
        "若只在分歧区变差，结论应聚焦评价参考风险；若一致区也差，模型部署问题更可信。",
    )

    h2(doc, "AlphaEarth 的 Conformal 结果如何读")
    prose(
        doc,
        "当前记录中，目标 90% coverage 的空间块测试边际 coverage 约 88.99%，平均 set size 约 1.449。"
        "这表示候选集合总体较精简，但在空间转移下出现轻微总体 coverage debt。"
        "下一步不是只盯 88.99%，而是检查 debt 是否集中在某些国家、类别或产品分歧区，并给出支持量与置信区间。",
    )

    page_break(doc)
    h1(doc, "第八部分｜高频导师追问与建议回答")
    h2(doc, "定位、理论与证据")

    qa(
        doc,
        "Q1：GeoBWER 是公平性指标，还是可靠性指标？",
        "它是以部署切片可靠性差距为构念的公平性审计指标。风险可以来自分类错误、分割损失、miscoverage 或漏检。"
        "它不声称覆盖公平性的所有社会维度，而是给 GeoFM/RSFM 提供跨任务、可操作的 deployment-slice fairness 入口。",
    )
    qa(
        doc,
        "Q2：它与 worst-group risk 或 CVaR 有什么不同？",
        "尾部风险的数学思想与 CVaR/upper-tail risk 有理论亲缘，但 GeoBWER 的贡献不只是一条聚合公式："
        "它把审计测度、切片语义、共同支持、空间依赖、标准化、选择性服务、conformal/CRC 和报告卡冻结为 GeoFM 协议。"
        "论文不应声称首次发明所有尾部风险数学，而应证明这是第一个成熟、跨任务、可复用的 GeoFM 部署切片公平审计框架。",
    )
    qa(
        doc,
        "Q3：为什么不用 accuracy、IoU 或最差组一个数就够了？",
        "accuracy/IoU 回答平均任务性能；最差组只看一个极端切片，容易受噪声影响。"
        "GeoBWER 看一个预注册的最差部署质量，并保留绝对平均、尾部风险和统计证书，因此信息层次不同。",
    )
    qa(
        doc,
        "Q4：为什么必须有三个随机种子？",
        "训练模型和 probe 都可能因初始化、批次顺序或抽样而变化。三个 seed 不能消除所有不确定性，"
        "但能识别排名或尾部身份是否只来自单次偶然运行。主结论应看 seed 汇总与配对区间。",
    )
    qa(
        doc,
        "Q5：为什么模型概率比只保存预测标签重要？",
        "Selective 需要置信度；LAC/APS/RAPS 需要所有类别概率；CRC 需要每标签或每像素概率；"
        "校准和 log loss 也需要概率。只保存 top-1 类别会使大部分不确定性扩展无法后处理。",
    )
    qa(
        doc,
        "Q6：Conformal 是不是 Selective-BWER 的升级版？",
        "不是替代关系。Selective 允许拒答，只评价被接受样本；Conformal 通常保留样本但扩大候选集合。"
        "一个研究服务覆盖，一个研究标签覆盖。二者可以并列，不能互相替代。",
    )
    qa(
        doc,
        "Q7：Conformal 会不会改变 Raw-GeoBWER 分数？",
        "不会。它读取同一概率资产，生成新的 miscoverage 风险表，再计算 Conformal-GeoBWER。"
        "Raw、Standardised 和 Selective 结果保持各自定义；这是后处理扩展，不需要重新训练基础模型。",
    )
    qa(
        doc,
        "Q8：为什么四个任务不能用完全相同的 conformal 算法？",
        "因为输出几何不同：单标签有唯一真类，适合 prediction sets；多标签本来就是集合，适合控制漏标比例；"
        "分割是高维像素场，像素相关且正负极不平衡，适合按图/事件控制风险。统一的是 calibration→risk→GeoBWER 接口，不是同一段集合代码。",
    )
    qa(
        doc,
        "Q9：GeoConformal 已经全部实现了吗？",
        "split CP/CRC 的任务适配器已经存在；地理核局部化被实现为有支持门控的 empirical comparator。"
        "不同任务能否形成正式结果取决于坐标、独立 calibration、局部有效样本量和概率资产。"
        "AlphaEarth 最适合；fMoW 次之；reBEN 与 Sen1 以 CRC 为正式主线，局部化只有支持通过才运行。",
    )
    qa(
        doc,
        "Q10：交叉切片支持少，是数据集问题还是门槛太严格？",
        "两者都可能影响，但不能用降低门槛来自动解决。高维交叉天然稀疏；大数据集也可能集中在少数组合。"
        "正确做法是预注册多分辨率轴、报告 support frontier、区分固定全集与 supported-universe，"
        "必要时用 Region×Superclass 做确认性轴，把 Country×原始类别保留为探索轴。",
    )
    qa(
        doc,
        "Q11：为什么支持不足还值得报告？",
        "因为“哪些部署区域不可审计”本身是数据覆盖结论，可直接指导数据采集。"
        "但论文至少需要一个支持密集任务给出非零 LCB 和较完整 coverage，AlphaEarth 正承担这一角色。",
    )
    qa(
        doc,
        "Q12：Standardised 会不会把真实差距洗掉？",
        "它回答的是条件化问题，不替代 Raw。Raw 描述实际部署风险，Standardised 描述在共同任务组成下的地理差距。"
        "两者一起报告：若 Raw 高、Standardised 低，组成解释较多；若两者都高，同类地理差距更可信。",
    )
    qa(
        doc,
        "Q13：空间 bootstrap 为什么必要？",
        "相邻像素、同一地点多时相图像和同一事件芯片共享环境因素。普通 i.i.d. bootstrap 会把它们当成许多独立证据，"
        "导致区间过窄。当前正式实现按 event/site/source tile/spatial block 构造 cluster influence，再以 multiplier bootstrap"
        "让同一 cluster 共同扰动；它不重新训练模型，也不改变 GeoBWER 点估计。",
    )
    qa(
        doc,
        "Q14：为什么叫 GeoBWER，而不是只叫 BWER2？",
        "GeoBWER 强调审计测度、空间依赖、地理部署单元和 GeoFM 适用性；BWER2 更像内部版本号。"
        "论文中可写 GeoBWER framework，metric version 则保留 geobwer_fractional_1.1 供复现。",
    )
    qa(
        doc,
        "Q15：这个框架以后别人怎么用？",
        "核心接口应接收逐样本预测/概率、真实标签、独立单元、切片和 cluster 元数据，而不要求接入模型权重。"
        "用户可通过 audit_multiclass、audit_multilabel、audit_segmentation、audit_conformal 或统一 CLI 生成 Audit Card。"
        "模型可以是开源 GeoFM、闭源 API 或普通 CNN。",
    )
    qa(
        doc,
        "Q16：当前最谨慎但有竞争力的论文主张是什么？",
        "GeoBWER 提供一套面向 GeoFM 的模型无关、任务感知、支持感知且兼顾空间依赖的部署切片公平审计协议；"
        "四任务套件显示平均性能、尾部风险、选择性服务和不确定性覆盖可以给出不同模型判断。"
        "是否能进一步使用“首个成熟可复用指标框架”的措辞，要以最终文献定位和完整正式结果为准。",
    )

    page_break(doc)
    h1(doc, "第九部分｜指标与缩写速查")
    table(
        doc,
        ["缩写/术语", "一句话解释"],
        [
            ["S1 / S2", "Sentinel-1 合成孔径雷达 / Sentinel-2 多光谱光学。"],
            ["SAR", "主动雷达成像，能穿云但有 speckle；S1 属于 SAR。"],
            ["IoU", "预测洪水区与真实洪水区交集 ÷ 并集；越高越好。"],
            ["Dice/F1", "预测与真实区域重叠程度；越高越好。"],
            ["BCE", "多标签每一类的二元交叉熵风险；越低越好。"],
            ["mAP / macro AP", "多标签各类平均 precision–recall 表现；macro 让类别等权。"],
            ["β", "预注册的最差部署质量比例，例如 10%。"],
            ["μ", "审计测度，即各切片在目标部署世界中的权重。"],
            ["support", "一个切片中可用于估计的独立样本/cluster/像素支持量。"],
            ["common support", "两个模型都有效覆盖、可配对比较的同一部署人口。"],
            ["CI / LCB", "置信区间 / 置信下界，用于区分点估计与可认证差距。"],
            ["coverage（Selective）", "被模型接受并自动处理的样本比例。"],
            ["coverage（Conformal）", "真实标签被预测集合包含的比例。"],
            ["set size", "单标签 conformal 每次输出的平均候选类别数。"],
            ["coverage debt", "实际 miscoverage 超过预设 α 的部分，或 coverage 低于目标的部分。"],
            ["LAC", "按类别概率阈值构造的简单 conformal 集合。"],
            ["APS", "按排序后的累计概率构造自适应集合。"],
            ["RAPS", "APS 加低排名类别正则，减少过大的集合。"],
            ["CRC", "Conformal Risk Control，用 calibration 控制一般单调风险的期望。"],
            ["GeoConformal comparator", "按地理邻近局部加权的经验空间不确定性比较器。"],
            ["protocol hash", "由冻结协议字段生成的指纹，用来阻止不同 estimand 的结果被误混。"],
        ],
        [1.75, 4.75],
    )

    h2(doc, "一句话区分最容易混的概念")
    prose(doc, "模型 vs embedding：模型是产生特征的参数系统；embedding 是模型产生的特征。", first_bold="模型 vs embedding：")
    prose(doc, "切片 vs cluster：切片是比较对象；cluster 是相关性与重采样单位。", first_bold="切片 vs cluster：")
    prose(doc, "Raw vs Standardised：Raw 是实际部署风险；Standardised 是共同组成下的条件化比较。", first_bold="Raw vs Standardised：")
    prose(doc, "Selective vs Conformal：Selective 允许拒答；Conformal 允许多答。", first_bold="Selective vs Conformal：")
    prose(doc, "CP vs CRC：CP 通常控制真实标签不在集合中的概率；CRC 控制更一般的平均风险。", first_bold="CP vs CRC：")
    prose(doc, "WorldCover vs Dynamic World：前者是主年度离散参考；后者是概率化近实时敏感性参考。", first_bold="WorldCover vs Dynamic World：")
    prose(doc, "数据集规模 vs 审计行数：一张多标签图会产生多条标签决策，但仍只有一个独立影像单元。", first_bold="数据集规模 vs 审计行数：")

    page_break(doc)
    h1(doc, "第十部分｜精选权威来源（用于追问时核对）")
    prose(
        doc,
        "以下只列与本次讲稿定义直接相关的原始论文或官方数据页。会议中无需逐条展示，"
        "但可用于回答“这个定义来自哪里”。",
    )
    source_line(doc, "Sen1Floods11（CVPRW 2020）", "https://openaccess.thecvf.com/content_CVPRW_2020/html/w11/Bonafilia_Sen1Floods11_A_Georeferenced_Dataset_to_Train_and_Test_Deep_Learning_CVPRW_2020_paper.html")
    source_line(doc, "reBEN / BigEarthNet v2 官方说明", "https://bigearth.net/")
    source_line(doc, "CROMA（NeurIPS 2023）", "https://papers.neurips.cc/paper_files/paper/2023/file/11822e84689e631615199db3b75cd0e4-Paper-Conference.pdf")
    source_line(doc, "TerraMind（ICCV 2025）", "https://openaccess.thecvf.com/content/ICCV2025/papers/Jakubik_TerraMind_Large-Scale_Generative_Multimodality_for_Earth_Observation_ICCV_2025_paper.pdf")
    source_line(doc, "DOFA", "https://arxiv.org/abs/2403.15356")
    source_line(doc, "Prithvi-EO-2.0", "https://arxiv.org/abs/2412.02732")
    source_line(doc, "AlphaEarth Satellite Embedding V1 官方数据页", "https://developers.google.com/earth-engine/datasets/catalog/GOOGLE_SATELLITE_EMBEDDING_V1_ANNUAL")
    source_line(doc, "ESA WorldCover 官方说明", "https://esa-worldcover.org/en/about/about")
    source_line(doc, "Dynamic World V1 官方数据页", "https://developers.google.com/earth-engine/datasets/catalog/GOOGLE_DYNAMICWORLD_V1")
    source_line(doc, "RAPS / prediction sets", "https://arxiv.org/abs/2009.14193")
    source_line(doc, "Conformal Risk Control（ICLR 2024）", "https://openreview.net/forum?id=33XGfHLtZg")
    source_line(doc, "GeoConformal Prediction", "https://doi.org/10.1080/24694452.2025.2516091")

    note(
        doc,
        "最后提醒：",
        "汇报的目标不是证明每个技术分支都已经完成，而是让导师看到：核心指标从直观分数升级为可复核协议，"
        "实验矩阵开始能回答机制问题，AlphaEarth 与 Conformal 补上了全球空间与不确定性两条关键证据线。",
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
